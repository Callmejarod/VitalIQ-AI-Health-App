"""Shared styling + helper functions for building the VitalIQ SRS docx."""
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

NAVY = RGBColor(0x1B, 0x3A, 0x57)
BLUE = RGBColor(0x2E, 0x6E, 0x9E)
TEAL = RGBColor(0x1F, 0x86, 0x8C)
GREY = RGBColor(0x44, 0x55, 0x66)
DARK = RGBColor(0x20, 0x2A, 0x33)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
AMBER = RGBColor(0xB9, 0x6A, 0x00)

DIAGRAMS = os.path.join(os.path.dirname(__file__), "diagrams")


def new_doc():
    doc = Document()
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = DARK
    pf = normal.paragraph_format
    pf.space_after = Pt(6)
    pf.line_spacing = 1.15
    for i, (size, color) in enumerate([(20, NAVY), (15, BLUE), (12.5, TEAL)], start=1):
        st = styles[f"Heading {i}"]
        st.font.name = "Calibri"
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = color
        st.paragraph_format.space_before = Pt(14 if i == 1 else 10)
        st.paragraph_format.space_after = Pt(4)
        st.paragraph_format.keep_with_next = True
    return doc


def set_margins(doc, inch=1.0):
    for s in doc.sections:
        s.top_margin = s.bottom_margin = Inches(inch)
        s.left_margin = s.right_margin = Inches(inch)


def shade(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    sh = OxmlElement("w:shd")
    sh.set(qn("w:val"), "clear")
    sh.set(qn("w:fill"), hexcolor)
    tcPr.append(sh)


def set_cell_text(cell, text, bold=False, color=None, size=10, align=None):
    cell.text = ""
    p = cell.paragraphs[0]
    if align:
        p.alignment = align
    run = p.add_run(str(text))
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.name = "Calibri"
    if color:
        run.font.color.rgb = color


def add_table(doc, headers, rows, widths=None, header_fill="1B3A57",
              zebra="EAF2F8", font=9.5):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        shade(hdr[i], header_fill)
        set_cell_text(hdr[i], h, bold=True, color=WHITE, size=font)
    for r_idx, row in enumerate(rows):
        cells = t.add_row().cells
        for i, val in enumerate(row):
            set_cell_text(cells[i], val, size=font)
            if zebra and r_idx % 2 == 1:
                shade(cells[i], zebra)
    if widths:
        for row in t.rows:
            for i, w in enumerate(widths):
                row.cells[i].width = Inches(w)
    doc.add_paragraph()
    return t


def H1(doc, text):
    return doc.add_heading(text, level=1)


def H2(doc, text):
    return doc.add_heading(text, level=2)


def H3(doc, text):
    return doc.add_heading(text, level=3)


def P(doc, text, bold=False, italic=False, size=11, color=None, align=None, space_after=None):
    p = doc.add_paragraph()
    if align:
        p.alignment = align
    if space_after is not None:
        p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    return p


def bullet(doc, text, level=0, bold_lead=None):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.3 + 0.3 * level)
    p.paragraph_format.space_after = Pt(2)
    if bold_lead:
        r = p.add_run(bold_lead)
        r.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p


def numbered(doc, text, bold_lead=None):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(2)
    if bold_lead:
        r = p.add_run(bold_lead); r.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p


def code_block(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.2)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
    # light shading
    pPr = p._p.get_or_add_pPr()
    sh = OxmlElement("w:shd")
    sh.set(qn("w:val"), "clear"); sh.set(qn("w:fill"), "F2F4F6")
    pPr.append(sh)
    return p


def add_image(doc, filename, width=6.3, caption=None):
    path = os.path.join(DIAGRAMS, filename)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(path, width=Inches(width))
    if caption:
        c = doc.add_paragraph()
        c.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = c.add_run(caption)
        r.italic = True
        r.font.size = Pt(9)
        r.font.color.rgb = GREY


def page_break(doc):
    doc.add_page_break()


def callout(doc, text, fill="EAF2F8"):
    t = doc.add_table(rows=1, cols=1)
    t.style = "Table Grid"
    cell = t.cell(0, 0)
    shade(cell, fill)
    set_cell_text(cell, text, size=10.5, color=NAVY)
    doc.add_paragraph()


def add_toc_field(doc):
    """Insert a real Word TOC field (updates on open / F9)."""
    p = doc.add_paragraph()
    run = p.add_run()
    fldBegin = OxmlElement("w:fldChar"); fldBegin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-3" \\h \\z \\u'
    fldSep = OxmlElement("w:fldChar"); fldSep.set(qn("w:fldCharType"), "separate")
    t = OxmlElement("w:t"); t.text = "Right-click and choose \u201cUpdate Field\u201d to build the Table of Contents."
    fldEnd = OxmlElement("w:fldChar"); fldEnd.set(qn("w:fldCharType"), "end")
    for el in (fldBegin, instr, fldSep, t, fldEnd):
        run._r.append(el)
