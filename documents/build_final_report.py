# -*- coding: utf-8 -*-
"""Build the VitalIQ FINAL REPORT (end-of-quarter deliverable).

This is NOT a proposal. It is a record of what was actually built, tested, and
shipped at submission time. The architecture described here matches the source
tree exactly:

  * The network/API layer is fully implemented: Retrofit talks to a working
    FastAPI + MongoDB backend with real AI insights.
  * The Repository layer is the SOLE data-access seam. Every ViewModel depends
    only on a repository INTERFACE; no ViewModel references Retrofit, Room, a
    DAO, or AppDatabase. Data-source wiring lives in one plain-Kotlin
    composition root (ServiceLocator), so swapping a data source changes only a
    RepositoryImpl and one wiring line — never a ViewModel.
  * Room provides on-device persistence behind the repositories as a
    write-through / offline-fallback cache (7 entities, 7 DAOs).
  * Android verification is by structured manual UAT; the backend has a real
    pytest suite.

Run:  python build_final_report.py
Out:  BTJ-FinalReport-document-06-24-26-brandonGalli.docx
"""
import io, os, sys
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8")

from docx.enum.text import WD_ALIGN_PARAGRAPH as ALIGN
from docx.shared import Pt, Inches

from docx_helpers import (
    new_doc, set_margins, add_table, H1, H2, H3, P, bullet, numbered,
    code_block, add_image, page_break, callout, add_toc_field,
    NAVY, BLUE, TEAL, GREY, DARK, WHITE, AMBER,
)

DOCDIR = os.path.dirname(__file__)
OUT = os.path.join(DOCDIR, "BTJ-FinalReport-document-06-24-26-brandonGalli.docx")

doc = new_doc()
set_margins(doc, 1.0)

C = ALIGN.CENTER

# ===========================================================================
# COVER PAGE
# ===========================================================================
P(doc, "VitalIQ", bold=True, size=40, color=NAVY, align=C, space_after=0)
P(doc, "AI-Assisted Personal Health & Fitness Companion", bold=True, size=15,
  color=BLUE, align=C, space_after=18)
P(doc, "FINAL PROJECT REPORT", bold=True, size=20, color=TEAL, align=C, space_after=4)
P(doc, "A record of what was built, tested, and delivered", italic=True, size=12,
  color=GREY, align=C, space_after=24)

add_table(doc,
    ["Field", "Detail"],
    [["Project", "VitalIQ — Android health & fitness companion with AI insights"],
     ["Team", "Team BTJ"],
     ["Author", "Brandon Galli"],
     ["Course", "DEV322 — Software Engineering"],
     ["Document", "Final Report (end-of-quarter)"],
     ["Version", "1.1 — Final (architecture-reconciled)"],
     ["Date", "June 24, 2026"],
     ["Repository", "VitalIQ (Android client + FastAPI backend)"],
     ["Build status", "Compiles & runs; full layered architecture integrated end-to-end"]],
    widths=[1.7, 5.0], font=10)

P(doc, "This report reflects the reality of the project at the time of submission. "
       "Every architectural claim below is verifiable in the submitted source tree.",
  italic=True, size=10, color=GREY, align=C)
page_break(doc)

# ===========================================================================
# DOCUMENT CONTROL
# ===========================================================================
H1(doc, "Document Control")

H2(doc, "Revision History")
add_table(doc,
    ["Version", "Milestone", "Date", "Summary"],
    [["1.0–6.0", "GP1–GP6 (SRS)", "May–Jun 2026",
      "Cumulative SRS: scope, user stories, use cases, diagrams, repository/API "
      "design, reactive-state architecture, test plans, and architectural audit."],
     ["Final 1.0", "Final Report", "Jun 24, 2026",
      "End-of-quarter record of the as-built system: layered Compose/MVVM client "
      "(UI → ViewModel → Repository → Room/network), AI insights, testing results, "
      "traceability, project-management summary, future work, and reflection."],
     ["Final 1.1", "Architecture reconciliation", "Jun 26, 2026",
      "Reconciled the report with the completed Repository-mediated architecture: "
      "the Repository layer is the sole data-access seam, data-source wiring is "
      "centralized in a plain-Kotlin composition root (ServiceLocator), and Room "
      "provides offline persistence behind the repositories."]],
    widths=[1.0, 1.3, 1.1, 3.9], font=9)

H2(doc, "Architecture Reconciliation Summary")
P(doc, "The GP5/GP6 SRS specified a Repository-mediated architecture with Room "
       "persistence behind the repositories. The final build implements that design "
       "in full. The table below maps each layer the SRS called for to where it lives "
       "in the submitted code.")
add_table(doc,
    ["Topic", "SRS design", "As-Built reality"],
    [["Network / API layer",
      "Retrofit ApiService consumed through the Repository.",
      "DONE. Retrofit + Gson ApiService (suspend functions) calls a running FastAPI "
      "backend; it is referenced only inside the RepositoryImpl classes."],
     ["Backend + AI insights",
      "External AI model consumed at runtime.",
      "DONE. FastAPI + MongoDB; insights generated by OpenAI gpt-4o-mini with a "
      "deterministic fallback scorer when the API is unavailable."],
     ["Repository mediator layer",
      "Repository is the sole data-access seam; ViewModels are data-source agnostic.",
      "DONE. Five repository interfaces + impls (Dashboard/Workout/Health/Insights/"
      "Profile). Every ViewModel depends only on an interface and calls repo methods; "
      "no ViewModel imports Retrofit, Room, a DAO, or AppDatabase."],
     ["Composition root (manual DI)",
      "Wiring centralized without a DI framework.",
      "DONE. com.vitaliq.app.di.ServiceLocator (plain Kotlin object, initialized in "
      "the VitalIQApp Application) builds AppDatabase + RetrofitClient.apiService once "
      "and exposes the five repository interfaces."],
     ["Local persistence / offline",
      "Room entities, DAOs, write-through cache behind the Repository.",
      "DONE. Room database (version 2) with 7 entities + 7 DAOs. Repositories write "
      "API results through to Room and fall back to the Room cache when the network "
      "is unavailable, so the app degrades to offline data instead of failing."],
     ["Automated Android tests",
      "Unit/acceptance test cases (VT/AT).",
      "Android verification is by structured manual UAT (documented in §5). The "
      "BACKEND has a real automated pytest suite (17 tests)."]],
    widths=[1.5, 2.6, 2.6], font=8.5)
callout(doc, "Integrity note: every “As-Built” statement in this report was verified "
             "against the source tree at submission. The instructor's architectural test — "
             "“if Retrofit were replaced by Room, only the Repository should change” — holds: "
             "because ViewModels depend on repository interfaces and all wiring is in the "
             "ServiceLocator, swapping a data source touches only a RepositoryImpl and one "
             "wiring line, never a ViewModel.")
page_break(doc)

# ===========================================================================
# TABLE OF CONTENTS
# ===========================================================================
H1(doc, "Table of Contents")
add_toc_field(doc)
page_break(doc)

# ===========================================================================
# 1. EXECUTIVE OVERVIEW
# ===========================================================================
H1(doc, "1. Executive Overview")

H2(doc, "1.1 Purpose")
P(doc, "VitalIQ is an Android application that helps everyday people understand their health "
       "by bringing fitness activity, vital-sign logging, and AI-generated guidance into one "
       "place. A companion backend stores the user’s data and produces a daily health score "
       "with personalized, data-cited suggestions. The goal was to deliver a working, "
       "architecturally disciplined MVP — not a clinical product — that demonstrates modern "
       "Android (Jetpack Compose + MVVM + coroutines + Repository + Room) talking to a real service.")

H2(doc, "1.2 The Problem We Set Out to Solve")
P(doc, "Health data is fragmented. Steps live in one app, blood pressure in a notebook, sleep "
       "in another tracker, and none of it is interpreted in plain language. People end up with "
       "numbers but no understanding. VitalIQ’s thesis: consolidate the daily signals a person "
       "can realistically capture on a phone, then use an AI model to translate that data into a "
       "single score and a short list of concrete, prioritized actions.")

H2(doc, "1.3 Target Users")
bullet(doc, "Health-conscious adults who already track some metrics and want them unified and interpreted.",
       bold_lead="Primary — ")
bullet(doc, "People managing a specific metric (e.g., blood pressure or weight) who want trend visibility and reminders.",
       bold_lead="Secondary — ")
bullet(doc, "Casual users who want a low-friction daily dashboard and an at-a-glance “how am I doing?” score.",
       bold_lead="Tertiary — ")

H2(doc, "1.4 What Was Built (Delivered)")
add_table(doc,
    ["Capability", "Delivered Behavior"],
    [["Daily dashboard", "Live summary (steps, workout minutes, sleep, water, HR, BP, health score) loaded through the Repository in parallel coroutines."],
     ["Workout tracking", "Start/stop sessions with a live timer, real accelerometer + step-counter sensors, activity classification, and intensity; session submitted via the Repository on stop."],
     ["Health & lifestyle logging", "Weight, blood pressure, heart rate, medication, meals, water, and sleep — each persisted through the Repository (API + Room) with validation."],
     ["AI Insights", "On-demand health score (0–100), six category scores, and ranked suggestions generated by OpenAI gpt-4o-mini, with a deterministic fallback."],
     ["History", "Lists of past workouts and biometric entries (weight, BP) retrieved through the Repository, served from Room when offline."],
     ["Profile", "Create/update profile (age, height, weight, goals) and import history from JSON/CSV; the ViewModel passes neutral file bytes and the Repository owns the upload transport."],
     ["Repository + persistence", "Five repository interfaces mediate all data; Room (7 entities/DAOs) caches results for offline use behind those repositories."],
     ["Navigation", "Full Jetpack Navigation: NavHost + bottom navigation across Home, Workout, Log, Insights, Profile, plus a History route; back stack survives rotation."],
     ["Backend service", "FastAPI + MongoDB with ~20 REST endpoints, file import, dashboard aggregation, and AI insight generation/history."]],
    widths=[1.9, 4.8], font=9)

H2(doc, "1.5 Scope Boundaries (Out of MVP Scope)")
P(doc, "Stated plainly, in the spirit of an accurate record, the following were intentionally "
       "outside the MVP scope:")
bullet(doc, "a single hardcoded “default” user; no login, accounts, or per-user isolation.",
       bold_lead="Authentication / multi-user — ")
bullet(doc, "the Android module uses structured manual UAT rather than an automated client test suite. (The backend does have automated tests.)",
       bold_lead="Automated Android tests — ")
bullet(doc, "history is loaded in full rather than paginated; large datasets are not yet chunked.",
       bold_lead="Pagination — ")

H2(doc, "1.6 Status at Submission")
callout(doc, "VitalIQ is a working, end-to-end MVP with the full layered architecture in place: "
             "the Compose UI, MVVM/coroutine spine, Repository seam, and Room persistence are "
             "integrated with a live FastAPI + MongoDB + AI backend, and every core user story is "
             "demonstrable — including rotation survival and graceful offline/failure handling. "
             "Remaining items (authentication, automated client tests, pagination) are scoped as "
             "future work, not architectural gaps.")
page_break(doc)

# ===========================================================================
# 2. FINAL SRS (REVISED)
# ===========================================================================
H1(doc, "2. Final Software Requirements Specification (Revised)")
P(doc, "This is the final, reconciled SRS. Requirements are marked Met, Partially Met, or "
       "Not Met against the shipped build.")

H2(doc, "2.1 Functional Requirements — Final Status")
add_table(doc,
    ["ID", "Requirement", "Status"],
    [["FR-1", "User can start and stop a workout and see a live timer.", "Met"],
     ["FR-2", "App tracks steps and classifies activity (stationary/walking/mixed/running) during a workout.", "Met (real accelerometer + step counter; simulation is opt-in for emulators)"],
     ["FR-3", "Completed workouts are submitted to and stored by the backend.", "Met (via WorkoutRepository)"],
     ["FR-4", "User can log blood pressure (systolic/diastolic) with validation.", "Met"],
     ["FR-5", "User can log weight, heart rate, medication, meals, water, and sleep.", "Met"],
     ["FR-6", "User can generate an AI health score with category scores and suggestions.", "Met (gpt-4o-mini; deterministic fallback on failure)"],
     ["FR-7", "User can view a daily dashboard summarizing today’s metrics.", "Met"],
     ["FR-8", "User can view history of past workouts and biometric entries.", "Met"],
     ["FR-9", "User can create/update a profile (age, height, weight, goals).", "Met"],
     ["FR-10", "User can import historical data from JSON/CSV.", "Met (Repository-mediated multipart upload)"],
     ["FR-11", "App caches data locally for offline use.", "Met (Room write-through cache + offline fallback in every repository)"],
     ["FR-12", "Data access mediated by a Repository layer.", "Met (5 repository interfaces; ViewModels are data-source agnostic)"],
     ["FR-13", "Per-user accounts and authentication.", "Not Met (single default user — out of MVP scope)"]],
    widths=[0.6, 4.7, 1.4], font=9)

H2(doc, "2.2 Non-Functional Requirements — Final Status")
add_table(doc,
    ["ID", "Requirement", "Status / Evidence"],
    [["NFR-1", "Reactive UI: state survives configuration changes (rotation).", "Met — ViewModel + StateFlow retained across rotation; nav back stack via rememberNavController."],
     ["NFR-2", "No blocking of the main thread; network/IO off the UI thread.", "Met — repositories wrap all IO in withContext(Dispatchers.IO); ViewModels use viewModelScope."],
     ["NFR-3", "No GlobalScope / runBlocking in the codebase.", "Met — static scan: none found; all launches in viewModelScope."],
     ["NFR-4", "Clear layering; data access only through the Repository.", "Met — UI → ViewModel → Repository → (Retrofit | Room). No ViewModel imports Retrofit/Room/DAO; wiring is centralized in ServiceLocator."],
     ["NFR-5", "Graceful handling of backend/AI failures.", "Met — ViewModels expose Error states with Retry; repositories fall back to the Room cache; backend falls back to deterministic scoring."],
     ["NFR-6", "Responsive: load operations run in parallel where possible.", "Met — dashboard/history use async/await + supervisorScope."]],
    widths=[0.6, 3.4, 2.7], font=9)

H2(doc, "2.3 Assumptions & Scope")
bullet(doc, "All data is mediated by the Repository layer; Room is the on-device cache and the API is the system of record. Repositories prefer the network and fall back to Room when offline.")
bullet(doc, "The runtime AI is OpenAI gpt-4o-mini, called server-side, with a deterministic fallback so the feature degrades instead of failing.")
bullet(doc, "Navigation is a full bottom-nav + NavHost graph whose state survives configuration changes.")
bullet(doc, "Scope held to a single default user; accounts and security were explicitly out of MVP scope.")

H2(doc, "2.4 Final User Stories")
add_table(doc,
    ["ID", "As a…", "I want to…", "So that…"],
    [["US-1", "active user", "start a workout and see live steps/activity", "I can track a session as it happens."],
     ["US-2", "health tracker", "log my blood pressure and weight", "I can monitor vitals over time."],
     ["US-3", "user seeking guidance", "generate an AI health score with suggestions", "I understand my data and what to improve."],
     ["US-4", "new user", "create a profile and import past data", "the app reflects me and my history."],
     ["US-5", "returning user", "see a daily dashboard", "I get an at-a-glance view of today."]],
    widths=[0.6, 1.5, 2.6, 2.0], font=9)

H2(doc, "2.5 Use Cases (Final)")
P(doc, "The four core use cases below were realized end-to-end. Diagrams appear in §4.")
add_table(doc,
    ["Use Case", "Actor", "Outcome (as-built)"],
    [["UC-1 Track Workout", "User", "Session recorded with sensor-driven steps/activity/intensity and persisted through the Repository."],
     ["UC-2 Log Blood Pressure", "User", "Validated reading stored via the Repository and visible in history."],
     ["UC-3 Generate AI Insights", "User", "Score + category scores + ranked suggestions returned (AI or fallback)."],
     ["UC-4 Manage Profile", "User", "Profile saved through the Repository; history importable from file."]],
    widths=[1.8, 1.0, 3.9], font=9)
page_break(doc)

# ===========================================================================
# 3. SYSTEM ARCHITECTURE (AS-BUILT)
# ===========================================================================
H1(doc, "3. System Architecture (As-Built)")
P(doc, "VitalIQ is a two-tier system: a Jetpack Compose Android client and a FastAPI backend "
       "with MongoDB. The client follows MVVM with a unidirectional reactive state flow and a "
       "Repository seam over two data sources (Retrofit network + Room cache). The description "
       "below is the architecture as actually implemented.")

H2(doc, "3.1 High-Level Topology")
code_block(doc,
    "Android client (Kotlin / Jetpack Compose)\n"
    "  UI (Composable screens)                         [Main]\n"
    "      |  collectAsStateWithLifecycle()\n"
    "  ViewModel (StateFlow<UiState>, viewModelScope)  [Main]\n"
    "      |  repo.<suspend fun>()        <-- depends on a Repository INTERFACE only\n"
    "  Repository (interface + Impl)                   [switches to IO]\n"
    "      |  withContext(Dispatchers.IO)\n"
    "      |-- Retrofit + Gson ApiService --HTTP/JSON--> FastAPI backend\n"
    "      |                                               |  Motor (async)\n"
    "      |                                           MongoDB (8 collections)\n"
    "      |                                               |  httpx\n"
    "      |                                           OpenAI gpt-4o-mini (insights) + fallback\n"
    "      '-- Room DAO/Entity (on-device cache) <-- write-through + offline fallback\n"
    "\n"
    "  ServiceLocator (plain-Kotlin composition root) builds ApiService + AppDatabase\n"
    "  once and hands each ViewModel its repository interface. No DI framework.")

H2(doc, "3.2 UI Layer (Jetpack Compose)")
P(doc, "Six Composable screens — Dashboard, Workout, Log, Insights, History, Profile — render "
       "purely from ViewModel state. Screens collect state with collectAsStateWithLifecycle() "
       "and trigger load() on resume. Reusable components (HealthRing, LineChart, StatCard, "
       "PrimaryButton, SectionHeader) keep screens declarative. MainActivity only hosts the "
       "Compose tree and theme — no business logic lives in the Activity or Composables.")

H2(doc, "3.3 ViewModel Layer")
P(doc, "Each screen has a dedicated ViewModel that owns a MutableStateFlow privately and exposes "
       "a read-only StateFlow publicly. State is modeled as sealed classes (Loading / Success / "
       "Error, plus domain-specific states such as the workout’s Idle / Active / Submitting / "
       "Summary). All asynchronous work is launched in viewModelScope and cancelled automatically "
       "in onCleared(). Critically, a ViewModel’s only data dependency is a repository INTERFACE — "
       "it has no knowledge of Retrofit, Room, DAOs, or AppDatabase.")
P(doc, "Representative ViewModel (Dashboard) — repository-mediated, explicit StateFlow, parallel loads:",
  bold=True, size=10, color=BLUE)
code_block(doc,
    "class DashboardViewModel(\n"
    "    private val repo: DashboardRepository      // interface only — data-source agnostic\n"
    ") : ViewModel() {\n\n"
    "    private val _uiState = MutableStateFlow<DashboardUiState>(DashboardUiState.Loading)\n"
    "    val uiState: StateFlow<DashboardUiState> = _uiState        // explicit StateFlow\n\n"
    "    fun load() = viewModelScope.launch {                       // viewModelScope only\n"
    "        try {\n"
    "            supervisorScope {\n"
    "                val summary = async { repo.getDashboardSummary() }  // parallel, via repo\n"
    "                val profile = async { repo.getProfile() }\n"
    "                _uiState.value = Success(summary.await(), profile.await())\n"
    "            }\n"
    "        } catch (e: Exception) {\n"
    "            _uiState.value = Error(e.message ?: \"Failed to load dashboard\")\n"
    "        }\n"
    "    }\n"
    "}")

H2(doc, "3.4 Data Access — Repository Seam (Composition Root)")
P(doc, "The Repository layer mediates between ViewModels and data sources (Retrofit network + "
       "Room cache). Each repository is an interface with a single implementation; the ViewModel "
       "receives the interface and never sees a concrete source. All data-source construction is "
       "centralized in one plain-Kotlin composition root — ServiceLocator — initialized once by "
       "the VitalIQApp Application. This is the manual-DI approach the course allows (no Hilt/"
       "Dagger/Koin).")
code_block(doc,
    "// di/ServiceLocator.kt — the ONLY place that knows the concrete sources\n"
    "object ServiceLocator {\n"
    "    private val database by lazy { AppDatabase.getInstance(appContext) }\n"
    "    private val api by lazy { RetrofitClient.apiService }\n"
    "    val dashboardRepository: DashboardRepository by lazy {\n"
    "        DashboardRepositoryImpl(api, database.profileDao())\n"
    "    }\n"
    "    // ...workout / health / profile / insights repositories\n"
    "}\n\n"
    "// A ViewModel factory pulls the ready-made interface — no Retrofit/Room here:\n"
    "val Factory = viewModelFactory {\n"
    "    initializer { DashboardViewModel(ServiceLocator.dashboardRepository) }\n"
    "}")
P(doc, "Each RepositoryImpl performs the network call and writes the result through to Room, then "
       "falls back to the Room cache on failure — so reads survive a dropped connection:")
code_block(doc,
    "override suspend fun getProfile(): ProfileDto = withContext(Dispatchers.IO) {\n"
    "    try {\n"
    "        val dto = api.getProfile()         // network (system of record)\n"
    "        profileDao.insert(dto.toEntity())  // write-through to Room\n"
    "        dto\n"
    "    } catch (e: Exception) {\n"
    "        profileDao.get()?.toDto() ?: throw e   // offline fallback from Room\n"
    "    }\n"
    "}")
add_table(doc,
    ["Aspect", "How the Repository seam delivers it"],
    [["Mediation", "ViewModels depend on interfaces; the Repository hides whether data came from Retrofit or Room."],
     ["Caching", "Write-through Room cache; reads fall back to Room when the network is unavailable."],
     ["Testability", "A ViewModel can be unit-tested against a fake repository interface — no HTTP/DB needed."],
     ["Swap sources", "Replacing Retrofit with Room edits only the RepositoryImpl + one ServiceLocator line; no ViewModel changes."]],
    widths=[1.3, 5.4], font=9)

H2(doc, "3.5 Local Datastore (Room)")
P(doc, "Persistence is a Room database (AppDatabase, version 2) with seven entities (workouts, "
       "health entries, profile, insights, medications, nutrition, sleep) and seven matching DAOs. "
       "Structured fields (e.g. a workout’s activity breakdown, a health entry’s value map) are "
       "stored as JSON String columns, so no Room TypeConverter is required. Room’s annotation "
       "processor runs via KSP. The database is reached only through the repositories, on "
       "Dispatchers.IO; the singleton is built with the application context to avoid leaks. The "
       "practical effect is offline resilience: previously loaded data remains visible when the "
       "backend is unreachable.")

H2(doc, "3.6 Backend Service")
P(doc, "A FastAPI application exposes ~20 REST endpoints under /api, persisted to MongoDB via the "
       "async Motor driver across eight collections (profiles, workouts, health_entries, "
       "medications, nutrition, sleep, ai_insights, imported_records). Pydantic models validate "
       "all input.")
add_table(doc,
    ["Endpoint group", "Examples", "Purpose"],
    [["Profile", "GET/PUT /profile", "Read/update the (seeded) user profile."],
     ["Workouts", "POST/GET /workouts", "Store and list sessions (sorted by start time)."],
     ["Health entries", "POST/GET/DELETE /health-entries", "Weight, BP, HR, body fat; filter by type."],
     ["Lifestyle", "/medications, /nutrition, /sleep", "Medication, meals/water, and sleep logs."],
     ["Dashboard", "GET /dashboard/summary", "Aggregates today’s metrics + health score."],
     ["Insights", "POST /insights/generate, GET /insights/latest|history", "AI scoring with deterministic fallback."],
     ["Import", "POST /import", "JSON/CSV bulk import of historical records."]],
    widths=[1.4, 2.5, 2.8], font=9)
P(doc, "AI insights: the backend builds a structured prompt from recent user data and calls "
       "OpenAI gpt-4o-mini (temperature 0.3) for a JSON response with overall_score, six "
       "category_scores, ranked suggestions, and a trend summary. If the key is missing or the "
       "call fails, a deterministic fallback scorer computes category scores from the same data "
       "and flags the result with used_fallback=true — so the feature degrades gracefully "
       "instead of erroring. The health score is a transparent heuristic, not a medical algorithm.")

H2(doc, "3.7 End-to-End Data Flow (Generate Insights)")
numbered(doc, "User taps Generate on the Insights screen; the screen calls InsightsViewModel.generateInsights().")
numbered(doc, "The ViewModel emits Loading, then launches in viewModelScope and calls repo.generateInsights() (suspend) — it does not know whether the result comes from the network or the cache.")
numbered(doc, "InsightsRepositoryImpl switches to Dispatchers.IO, calls the API, writes the result through to Room, and returns it (falling back to the cached insight on failure).")
numbered(doc, "Retrofit issues an HTTP POST to the FastAPI backend off the main thread; the backend calls gpt-4o-mini (or falls back) and returns the insight.")
numbered(doc, "The ViewModel maps the response into Success(state) (or Error) on the StateFlow.")
numbered(doc, "Compose recomposes from the new state, rendering the score, category scores, and suggestions.")

H2(doc, "3.8 Concurrency, Threading & Lifecycle Discipline")
add_table(doc,
    ["Discipline", "How it was maintained"],
    [["Scope ownership", "Every coroutine launches in viewModelScope; cancelled on onCleared()."],
     ["Thread boundary", "ViewModels/state live on Main; the Repository switches to Dispatchers.IO for every network + Room call (withContext(Dispatchers.IO))."],
     ["No GlobalScope/runBlocking", "Static scan across the codebase: zero occurrences."],
     ["Main never blocks", "All network and disk calls are suspend functions executed off the UI thread on the IO dispatcher."],
     ["Config-change survival", "ViewModel + StateFlow retain state across rotation; nav back stack via rememberNavController (verified manually)."],
     ["Sensor cleanup", "WorkoutViewModel unregisters its SensorEventListener and cancels the timer job in onCleared()."],
     ["Lifecycle logging", "A single 'VitalIQ' log tag across Application, MainActivity (onCreate/onResume/onPause/onStop/onDestroy), and every ViewModel (init/onCleared)."]],
    widths=[1.9, 4.8], font=9)

H2(doc, "3.9 Architectural Discipline — Verdict")
callout(doc, "The full layered architecture is in place and disciplined: the UI is a pure function "
             "of ViewModel state; ViewModels own state and threading and depend only on repository "
             "interfaces; the Repository is the sole data-access seam over Retrofit + Room; and all "
             "IO runs on Dispatchers.IO. Coroutine and lifecycle hygiene are intact, and the "
             "data-source wiring is isolated in one composition root, so the architecture satisfies "
             "the course’s separation-of-layers requirement end to end.")
page_break(doc)

# ===========================================================================
# 4. DESIGN ARTIFACTS
# ===========================================================================
H1(doc, "4. Design Artifacts (Diagrams, Wireframes, Screens)")
P(doc, "The diagrams below were authored during design and remain accurate for the four core "
       "use cases and the implemented architecture. Note: live device screenshots are not embedded "
       "in this document; instead, each final screen is described against the implemented UI so the "
       "record matches the shipped build.")

H2(doc, "4.1 Use-Case Diagrams")
add_image(doc, "uc1_workout.png", width=5.6, caption="UC-1 Track Workout — actor, system, and core interactions.")
add_image(doc, "uc2_bp.png", width=5.6, caption="UC-2 Log Blood Pressure — validated vital-sign entry.")
add_image(doc, "uc3_ai.png", width=5.6, caption="UC-3 Generate AI Insights — on-demand scoring and suggestions.")
add_image(doc, "uc4_profile.png", width=5.6, caption="UC-4 Manage Profile — profile creation and history import.")
page_break(doc)

H2(doc, "4.2 Activity Diagrams")
add_image(doc, "act_workout.png", width=6.2, caption="Workout activity flow (start → track → stop → submit).")
add_image(doc, "act_insights_week5.png", width=6.2, caption="AI Insights activity flow with Main/IO separation.")
page_break(doc)

H2(doc, "4.3 Sequence & State Diagrams")
add_image(doc, "seq_insights_week5.png", width=6.2, caption="Insights sequence: UI → ViewModel → Repository → network → backend → AI → state.")
add_image(doc, "state_workout.png", width=4.6, caption="Workout state machine (Idle → Active → Submitting → Summary).")
P(doc, "As-built note: the sequence diagram’s Repository hop matches the shipped code — the "
       "ViewModel calls the repository, which mediates the Retrofit call and the Room cache.",
  italic=True, size=9.5, color=GREY)
page_break(doc)

H2(doc, "4.4 Final Screen Inventory (Wireframe Description)")
add_table(doc,
    ["Screen", "Final Layout & Key Elements"],
    [["Dashboard (Home)", "Health ring (today’s score) atop stat cards: steps, workout minutes, sleep, water, HR, BP. Pulls from /dashboard/summary via the Repository."],
     ["Workout", "Large timer, live step count (step-counter sensor), current activity label (accelerometer), intensity indicator; Start/Stop control. Idle/Active/Summary states."],
     ["Log", "Expandable cards for weight, BP, HR, medication, meal, water, sleep; each validates and posts through the Repository on save with a snackbar confirmation."],
     ["Insights", "Generate button; on success shows the 0–100 score, six category sub-scores, a trend summary, and a ranked suggestion list."],
     ["History", "Scrollable lists of past workouts and biometric entries (weight, BP) with a small line chart for trends; served from Room when offline."],
     ["Profile", "Editable fields (name, age, height, weight, goal); Save persists via the Repository; Import pulls history from a JSON/CSV file."]],
    widths=[1.4, 5.3], font=9)
page_break(doc)

# ===========================================================================
# 5. TESTING DOCUMENTATION
# ===========================================================================
H1(doc, "5. Testing Documentation")
P(doc, "Testing combined an automated backend suite with structured manual validation of the "
       "Android client. This section reports what was tested, how, and the real results — "
       "including defects found and fixed, and the gaps that remain.")

H2(doc, "5.1 Test Strategy & Levels")
add_table(doc,
    ["Level", "Where", "Approach"],
    [["Unit", "Backend", "pytest against FastAPI endpoints (CRUD, validation, aggregation)."],
     ["Unit", "Android", "Structured manual verification of ViewModel/Repository behavior (see limitations)."],
     ["Integration", "Full stack", "Client → Repository → backend → MongoDB → AI exercised manually per feature."],
     ["Verification", "Client", "Manual checks that each unit behaves per the technical spec (state, threading, layering)."],
     ["Validation / UAT", "Client", "Manual walkthrough of each user story end-to-end on the emulator/device."],
     ["Regression", "Both", "Re-running backend suite + repeating core client flows after each change."]],
    widths=[1.4, 1.2, 4.1], font=9)

H2(doc, "5.2 Backend Automated Tests (pytest) — Real Results")
P(doc, "backend/tests/test_vitaliq_api.py contains 17 tests across the API surface. These run "
       "against a live backend + MongoDB; the AI test asserts a real model call (used_fallback "
       "is False) and therefore requires a valid key and network.")
add_table(doc,
    ["Suite", "Tests", "Covers", "Result"],
    [["Root", "1", "Health check returns ok.", "PASS"],
     ["Profile", "2", "Seed-on-first-read; update persistence.", "PASS"],
     ["Workouts", "1", "Create with sensor data; list sorted desc.", "PASS"],
     ["Health entries", "4", "Create weight/BP/HR; filter by type; delete.", "PASS"],
     ["Medications", "1", "Create, list, delete.", "PASS"],
     ["Nutrition", "2", "Meal + water; filter by kind.", "PASS"],
     ["Sleep", "1", "Create and list.", "PASS"],
     ["Dashboard", "1", "Summary contains all required keys.", "PASS"],
     ["Insights", "2", "Generate (real AI, not fallback); latest.", "PASS when key+network present"],
     ["Import", "2", "CSV and mixed-type JSON import counts.", "PASS"]],
    widths=[1.4, 0.7, 3.6, 1.6], font=9)

H2(doc, "5.3 Android Verification (Manual) — Technical Behavior")
add_table(doc,
    ["ID", "Unit Under Test", "Check", "Result"],
    [["VT-1", "WorkoutViewModel", "Start → Active state; timer increments each second.", "PASS"],
     ["VT-2", "StateFlow rotation", "Rotate during active workout; state retained.", "PASS"],
     ["VT-3", "Workout submit", "Stop → repo.createWorkout() called off main thread; UI not blocked.", "PASS"],
     ["VT-4", "Layer boundary", "No Retrofit/Room/DAO/AppDatabase import in any ViewModel (data goes through repository interfaces).", "PASS"],
     ["VT-5", "Coroutine hygiene", "Static scan for GlobalScope/runBlocking.", "PASS (none)"],
     ["VT-6", "LogViewModel validation", "Empty/invalid BP rejected before the repository call.", "PASS"],
     ["VT-7", "InsightsViewModel", "Loading emitted before the repository call.", "PASS"],
     ["VT-8", "Sensor classification", "Accelerometer magnitudes map to stationary/walking/mixed/running.", "PASS"],
     ["VT-9", "onCleared cleanup", "Timer cancelled; sensor listener unregistered.", "PASS"],
     ["VT-10", "Offline fallback", "With backend unreachable, repository serves cached Room data.", "PASS"]],
    widths=[0.5, 1.7, 3.5, 1.0], font=8.5)

H2(doc, "5.4 Validation / User Acceptance (Manual) — User Stories")
add_table(doc,
    ["ID", "Story", "User Steps", "Outcome", "Result"],
    [["AT-1", "US-1", "Open Workout; Start.", "Timer + live sensor-driven steps/activity shown.", "PASS"],
     ["AT-2", "US-1", "Start; rotate device.", "Session keeps running; metrics preserved.", "PASS"],
     ["AT-3", "US-2", "Enter valid BP; Save.", "Stored; appears in history.", "PASS"],
     ["AT-4", "US-2", "Enter invalid/empty BP; Save.", "Rejected with validation message.", "PASS"],
     ["AT-5", "US-3", "Open Insights; Generate.", "Score, category scores, suggestions shown.", "PASS"],
     ["AT-6", "US-3", "Generate with backend unreachable (airplane mode).", "Error state with Retry / cached fallback; no crash.", "PASS"],
     ["AT-7", "US-4", "Edit profile; Save.", "Persisted server-side; reloads on return.", "PASS"],
     ["AT-8", "US-5", "Open Home.", "Dashboard summary loads today’s metrics.", "PASS"]],
    widths=[0.5, 0.7, 2.2, 2.7, 0.7], font=8.5)

H2(doc, "5.5 Integration & Regression")
bullet(doc, "Integration: the full vertical slice (profile → workout → log → dashboard → insights) was exercised against the live backend through the repositories; data created on one screen appears on others after reload.")
bullet(doc, "Regression: the backend pytest suite was re-run after backend changes; core client flows were repeated after client changes. No regressions observed in core features.")
bullet(doc, "Sensors: real accelerometer + step counter drive the workout by default; SIMULATE_SENSORS=true is an opt-in flag for emulators that lack a step sensor.")

H2(doc, "5.6 Defects Found & Corrected")
add_table(doc,
    ["Defect", "Symptom", "Fix", "Re-test"],
    [["Step-counter base offset", "First reading showed device-lifetime steps.", "Capture baseStepCount on first event; report the delta.", "VT-8 PASS"],
     ["Insights loading flicker", "Brief empty flash before content.", "Emit Loading before the repository call.", "VT-7 PASS"],
     ["Rotation reset (early build)", "Workout reset on rotation in an early prototype.", "Move state into the ViewModel/StateFlow.", "VT-2 PASS"],
     ["Offline read failure (early build)", "Reads failed outright when the backend was down.", "Add Room write-through cache + offline fallback in the repositories.", "VT-10 PASS"]],
    widths=[1.7, 2.3, 2.3, 0.9], font=9)

H2(doc, "5.7 Results Summary")
add_table(doc,
    ["Metric", "Result"],
    [["Backend automated tests", "17 (16 deterministic PASS; 1 AI test PASS with key+network)"],
     ["Android verification checks (manual)", "10 / 10 PASS"],
     ["Android acceptance checks (manual)", "8 / 8 PASS"],
     ["Integration (core vertical slice)", "PASS"],
     ["Regression after changes", "No regressions in core features"],
     ["Rotation / config-change survival", "PASS"],
     ["Offline fallback (Room cache)", "PASS"],
     ["GlobalScope / runBlocking", "None found"]],
    widths=[4.0, 2.7], font=9.5)

H2(doc, "5.8 Known Limitations (Testing)")
bullet(doc, "Android client verification is manual rather than an automated JUnit/Espresso suite, so it is not reproducible in CI.")
bullet(doc, "The AI insights backend test is environment-dependent (needs a valid key and network); without them it fails by design rather than passing on the fallback.")
bullet(doc, "No automated UI (Compose/Espresso) tests; visual regressions would not be caught automatically.")
bullet(doc, "No load/performance or security testing was performed.")
page_break(doc)

# ===========================================================================
# 6. TRACEABILITY REVIEW
# ===========================================================================
H1(doc, "6. Traceability Review")
P(doc, "This review connects the final implementation back to the original goals: each user "
       "story traces through its primary requirements to the feature that delivers it and the "
       "test(s) that validate it.")
add_table(doc,
    ["User Story", "Requirements", "Implemented In", "Validated By", "Verdict"],
    [["US-1 Track workout", "FR-1, FR-2, FR-3", "WorkoutViewModel + sensors + WorkoutRepository", "VT-1–3, VT-8–9, AT-1–2", "Met"],
     ["US-2 Log vitals", "FR-4, FR-5", "LogViewModel + HealthRepository", "VT-6, AT-3–4", "Met"],
     ["US-3 AI insights", "FR-6", "InsightsViewModel + InsightsRepository + gpt-4o-mini", "VT-7, AT-5–6", "Met"],
     ["US-4 Profile + import", "FR-9, FR-10", "ProfileViewModel + ProfileRepository", "AT-7", "Met"],
     ["US-5 Dashboard", "FR-7, FR-8", "DashboardViewModel + DashboardRepository", "AT-8", "Met"],
     ["(cross-cutting)", "FR-11 offline cache", "Room + repository fallback", "VT-10, AT-6", "Met"],
     ["(cross-cutting)", "FR-12 Repository seam", "5 repository interfaces + ServiceLocator", "VT-4", "Met"],
     ["(cross-cutting)", "FR-13 auth", "— (out of MVP scope)", "—", "Not Met"]],
    widths=[1.5, 1.3, 2.1, 1.4, 0.8], font=8.5)
H2(doc, "6.1 Findings")
bullet(doc, "Coverage: all five core user stories are implemented and validated end-to-end; no core story is unverified.")
bullet(doc, "Cross-cutting architecture requirements (Repository seam, Room offline cache) are met and trace to concrete tests (VT-4, VT-10).")
bullet(doc, "The only Not-Met requirement (authentication) was explicitly out of MVP scope and is carried forward as future work.")
page_break(doc)

# ===========================================================================
# 7. PROJECT MANAGEMENT SUMMARY
# ===========================================================================
H1(doc, "7. Project Management Summary")

H2(doc, "7.1 Methodology & Cadence")
P(doc, "The team worked in weekly milestone increments (GP1–GP6) culminating in this final "
       "report, tracking work on a GitHub-based board and committing through feature branches "
       "with review before merge. Each milestone paired a documentation deliverable (SRS growth) "
       "with concrete implementation progress.")

H2(doc, "7.2 Milestone Log")
add_table(doc,
    ["Milestone", "Focus", "Outcome"],
    [["GP1–GP2", "Scope, requirements, user stories, narrative.", "Project scope and functional/non-functional requirements set."],
     ["GP3", "Use cases, diagrams, traceability, state machine.", "Design artifacts authored (use-case/activity/sequence/state)."],
     ["GP4", "Repository/API contract, DTOs, thread boundaries.", "Data + concurrency design specified."],
     ["GP5", "Reactive state architecture, wireframes, test plans.", "Compose UI + ViewModel/StateFlow spine built."],
     ["GP6", "Integration, testing, architectural audit.", "Network + Repository + Room wired to backend; core features integrated."],
     ["Final", "As-built record + architecture reconciliation.", "Working MVP delivered with the full layered architecture; report matches code."]],
    widths=[1.1, 2.7, 2.9], font=9)

H2(doc, "7.3 Project Board Summary")
add_table(doc,
    ["Column", "Representative Items at Close"],
    [["Done", "Compose screens; ViewModels/StateFlow; navigation; Repository layer + ServiceLocator; Room offline cache; Retrofit layer; FastAPI backend; AI insights + fallback; file import; backend pytest suite."],
     ["In progress", "Optional hardening (asStateFlow(), inline snackbar retry); pagination for large histories."],
     ["Backlog", "Authentication/multi-user; automated Android tests; charts/analytics polish; accessibility."]],
    widths=[1.4, 5.3], font=9)

H2(doc, "7.4 Team Retrospective")
add_table(doc,
    ["What went well", "What was hard", "What we’d carry forward"],
    [["Clean MVVM + StateFlow spine with a Repository seam; reactive UI came together quickly.",
      "Coordinating the Android client with a live backend + AI under a deadline.",
      "Keep the composition root (ServiceLocator) small and central so wiring stays in one place."],
     ["Real end-to-end integration (UI → Repository → backend → AI), not mocked.",
      "Emulators lack a step sensor, so sensor work had to be validated on a device.",
      "Stand up automated client tests earlier alongside the manual UAT."],
     ["Graceful degradation (AI fallback + Room offline cache) kept core features reliable.",
      "Keeping documentation in lockstep with fast-moving code.",
      "Continue reconciling the SRS against the code at each milestone."]],
    widths=[2.2, 2.2, 2.3], font=9)

H2(doc, "7.5 Challenges & How They Were Addressed")
bullet(doc, "AI reliability: wrapped the model call in a deterministic fallback so insights still return when the API is unavailable.", bold_lead="Challenge — ")
bullet(doc, "Offline resilience: added a Room write-through cache behind the repositories so reads survive a dropped connection.", bold_lead="Challenge — ")
bullet(doc, "Rotation resets in an early prototype: moved state ownership into the ViewModel + StateFlow.", bold_lead="Challenge — ")
bullet(doc, "Keeping ViewModels data-source agnostic: centralized all Retrofit/Room wiring in a ServiceLocator composition root so no ViewModel imports a data source.", bold_lead="Challenge — ")

H2(doc, "7.6 Lessons Learned")
bullet(doc, "A Repository seam plus a single composition root keeps data-source decisions in one place and makes ViewModels trivial to reason about and test.")
bullet(doc, "Manual testing scales poorly; even a thin automated client suite would make regression safer.")
bullet(doc, "Designing for graceful degradation (the AI fallback and the Room offline cache) paid off more than any single feature.")
bullet(doc, "Living documents drift; reconciling the SRS against the code at each milestone keeps the record trustworthy.")
page_break(doc)

# ===========================================================================
# 8. FUTURE DEVELOPMENT OPPORTUNITIES
# ===========================================================================
H1(doc, "8. Future Development Opportunities")
P(doc, "If development continued, the following work — ordered by priority — would close the gap "
       "between the current MVP and a robust product.")
add_table(doc,
    ["Area", "Improvement", "Why / Payoff"],
    [["Testing", "Add JVM ViewModel/Repository tests (with fake repositories) + Compose/Espresso UI tests in CI.", "Reproducible coverage; catches regressions automatically."],
     ["Security", "Add authentication and per-user data isolation; secret management for keys.", "Removes the single hardcoded user; protects data."],
     ["Persistence", "Add a real Room Migration path (replace fallbackToDestructiveMigration) and an observable Flow-based cache.", "Preserves cached data across schema changes; live-updating UI."],
     ["Performance", "Pagination for history; lazy lists; image/asset tuning.", "Smoother scrolling and lower latency on large histories."],
     ["AI quality", "Richer prompts, trend-aware scoring, and on-device summarization options.", "More accurate, explainable insights."],
     ["Accessibility", "Content descriptions, dynamic type, contrast and TalkBack passes.", "Usable by more people; meets accessibility expectations."],
     ["UX", "Charts/analytics polish, reminders/notifications, richer logging UIs.", "Higher engagement and daily value."],
     ["Resilience", "Surface a denied-permission message on the Workout screen; inline snackbar retry on Log failures.", "Clearer recovery paths for edge cases."]],
    widths=[1.2, 3.0, 2.5], font=9)
page_break(doc)

# ===========================================================================
# 9. REFLECTION & CONCLUSION
# ===========================================================================
H1(doc, "9. Reflection & Conclusion")

H2(doc, "9.1 What We Learned")
P(doc, "Building VitalIQ taught us that the hard part of modern app development is rarely a single "
       "feature — it is the disciplined wiring between layers and the decisions about what to "
       "build now versus later. We learned modern Android in depth: Jetpack Compose as a function "
       "of state, MVVM with StateFlow, a Repository seam over Retrofit + Room, and coroutine/"
       "lifecycle hygiene that survives rotation without leaks. We learned to integrate a real "
       "backend and a real AI model, and to design for failure so the product degrades gracefully "
       "instead of breaking.")

H2(doc, "9.2 On Architecture & Trade-offs")
P(doc, "Our most instructive work was getting the layering right: ViewModels that depend only on "
       "repository interfaces, a single composition root that owns data-source wiring, and a Room "
       "cache behind the repositories for offline resilience. The payoff is concrete — the "
       "instructor’s test (“replace Retrofit with Room and only the Repository should change”) "
       "holds, because no ViewModel knows where its data comes from. Keeping that seam clean made "
       "the rest of the system easier to reason about, test, and extend.")

H2(doc, "9.3 On Collaboration, Testing & Process")
P(doc, "Weekly milestones kept momentum and forced steady integration, but our reliance on manual "
       "client testing made regressions a manual risk; a small automated client suite would have "
       "repaid itself quickly. Keeping documentation honest — reconciling the SRS against the code "
       "at the end — turned out to be a core engineering activity, not an afterthought.")

H2(doc, "9.4 Conclusion")
callout(doc, "VitalIQ delivered a working, integrated MVP with a complete layered architecture: a "
             "reactive Compose client whose ViewModels own state and depend only on repository "
             "interfaces, a Repository seam over Retrofit + Room with offline fallback, and a live "
             "FastAPI + MongoDB + AI backend — with every core user story demonstrable end-to-end, "
             "including rotation survival and graceful failure handling. Remaining items "
             "(authentication, automated client tests, pagination) are scoped as future work rather "
             "than architectural gaps. We end the quarter with a product that works, an architecture "
             "that matches its documentation, and a clear path to grow.")

# ===========================================================================
# SAVE
# ===========================================================================
doc.save(OUT)
print("SAVED:", OUT)
print("paragraphs:", len(doc.paragraphs), "tables:", len(doc.tables),
      "images:", len(doc.inline_shapes))
