# -*- coding: utf-8 -*-
"""Weekly Team Status Report + Retrospective for VitalIQ (GP4 week)."""
import io, sys, os
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8")
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx_helpers import (new_doc, set_margins, add_table, H1, H2, P, bullet,
                          page_break, callout, NAVY, BLUE, GREY, TEAL)

doc = new_doc()
set_margins(doc, 1.0)

P(doc, "Weekly Team Status Report", bold=True, size=26, color=NAVY,
  align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
P(doc, "Including Team Retrospective", bold=True, size=13, color=BLUE,
  align=WD_ALIGN_PARAGRAPH.CENTER, space_after=14)
P(doc, "VitalIQ — AI-Powered Personal Health Intelligence Platform", size=13, color=TEAL,
  align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
P(doc, "Team BTJ  ·  Brandon Galli  ·  Jarod Atienzo  ·  Thaigo Amaro Da Silva", size=11,
  align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
P(doc, "Milestone: GP4 — Architectural Consolidation  ·  Week of Jun 3, 2026", size=11,
  italic=True, color=GREY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
P(doc, "GitHub: https://github.com/Callmejarod/VitalIQ-AI-Health-App", size=11, color=BLUE,
  align=WD_ALIGN_PARAGRAPH.CENTER)
doc.add_paragraph()

H1(doc, "1. Milestone Summary")
P(doc, "This week's milestone was a structural / architectural-design checkpoint. The goal was to "
  "consolidate the MVVM architectural spine and design how the existing Room-based Repository will "
  "evolve to support an external API as a second data source \u2014 without writing networking code "
  "and without breaking architectural boundaries. All planned deliverables were completed and "
  "incorporated into the cumulative SRS.")

H1(doc, "2. Deliverable Status")
add_table(doc, ["Deliverable", "Owner", "Status"],
    [["Repository Evolution Plan", "Brandon", "Complete"],
     ["API Contract Specification", "Jarod", "Complete"],
     ["DTO + Mapping Plan", "Thaigo", "Complete"],
     ["Thread Boundary Plan + coroutine write-up", "Brandon", "Complete"],
     ["Updated Activity Diagram (Main vs IO, dual source)", "Thaigo", "Complete"],
     ["Updated Sequence Diagram (dual source, scope/thread bands)", "Jarod", "Complete"],
     ["State Machine Diagram (carried + refined)", "Thaigo", "Complete"],
     ["Architectural Risk Statement", "Jarod", "Complete"],
     ["Cumulative SRS assembly + instructor-feedback corrections", "Brandon", "Complete"],
     ["Prior-week continuity (GP1\u2013GP3, Week 4)", "All", "Complete"]],
    widths=[3.8, 1.6, 1.2])

H1(doc, "3. Individual Contributions")
add_table(doc, ["Member", "Contributions This Week"],
    [["Brandon Galli", "Authored the Repository Evolution Plan and Thread Boundary Plan; wrote the coroutine-foundations section; assembled the cumulative SRS and applied instructor-feedback corrections."],
     ["Jarod Atienzo", "Authored the API Contract Specification from the backend endpoints; produced the updated sequence diagram; wrote the Architectural Risk Statement; maintained the GitHub repo and project board."],
     ["Thaigo Amaro Da Silva", "Authored the DTO + Mapping Plan; produced the updated activity diagram with Main/IO swimlanes and refined the state machine; reviewed layer-boundary compliance."]],
    widths=[1.6, 6.0])

H1(doc, "4. GitHub Project Board")
P(doc, "Project execution is tracked on the team GitHub Project board. This week's cards moved from "
  "\u201cTo Do\u201d through \u201cIn Progress\u201d to \u201cDone\u201d for each design deliverable; the "
  "API-contract and Repository-evolution cards were linked to the SRS sections they produced.")
bullet(doc, "Backlog refined with next-phase cards: Retrofit client setup, DTO classes, mapper implementation, Room CachedInsight entity + migration.")
bullet(doc, "No code-bearing cards were closed this week by design (architectural checkpoint only).")
page_break(doc)

H1(doc, "5. Team Retrospective")
H2(doc, "5.1 What Worked")
for b in ["Designing the Repository contract before any networking code clarified exactly which "
          "functions are local-only, remote-only, or cache-then-network. The branching-logic example "
          "removed ambiguity about where decisions live.",
          "Keeping three separate model families (DTO / Domain / Entity) on paper made the mapping "
          "responsibilities obvious and prevented the team from \u201cjust reusing the DTO everywhere.\u201d",
          "Drawing the Main\u2192IO\u2192Main bands on the sequence diagram made the thread-boundary plan "
          "concrete and caught two places where we had informally assumed mapping happened on Main.",
          "Splitting deliverables by member with a shared SRS outline avoided merge friction."]:
    bullet(doc, b)

H2(doc, "5.2 What Didn't Work")
for b in ["Our first activity diagram tried to put both the cache-read and API-call branches on the "
          "same row, which overlapped and was unreadable.",
          "We initially under-specified the error contract, so the risk statement had no concrete "
          "failure-handling story to point at.",
          "Early drafts let mapping drift toward the ViewModel in discussion, which would have violated "
          "the layer rules."]:
    bullet(doc, b)

H2(doc, "5.3 How We Fixed What Didn't Work")
for b in ["Re-laid out the activity diagram so each branch occupies its own row and merges before the "
          "Flow emission; thread-crossing edges are colored to read at a glance.",
          "Added a dedicated Error Response Format (\u00a716.4) with a status\u2192code\u2192handling table, "
          "which the concurrency/resilience risks now reference directly.",
          "Wrote the explicit rule that mapping lives only in the Repository (\u00a717.3) and added it to "
          "the structural-risk mitigations, then code-reviewed the contract against the Layer Boundary "
          "Rules."]:
    bullet(doc, b)

H2(doc, "5.4 Lessons Learned (to apply next time)")
for b in ["Define the API error envelope at the same time as the success schema \u2014 resilience design "
          "depends on it.",
          "Decide the dispatcher for every unit of work (including mapping) up front, and annotate it on "
          "the diagram rather than leaving it implicit.",
          "Lay out diagram swimlanes with one activity per row from the start; it mirrors the "
          "\u201cone activity per step\u201d pseudocode rule and avoids rework.",
          "Lock the Repository's public surface to domain types early so no later code accidentally "
          "leaks DTOs or entities upward."]:
    bullet(doc, b)

H1(doc, "6. Risks & Blockers")
add_table(doc, ["Item", "Impact", "Plan"],
    [["Backend JSON contract may still change before implementation.", "Medium", "Nullable DTO fields + lenient parsing absorb additive changes; mapper isolates breaking ones."],
     ["Emulator step-counter limitations for sensor testing.", "Low", "Accelerometer-derived step estimation fallback (already in scope)."],
     ["No blockers requiring instructor action this week.", "\u2014", "\u2014"]],
    widths=[3.6, 1.0, 3.0], font=9.5)

H1(doc, "7. Plan for Next Milestone")
for b in ["Implement the Retrofit client and ApiService against the documented contract.",
          "Add the CachedInsight Room entity, InsightDao, and a versioned migration.",
          "Implement DTO/Domain/Entity mappers exactly as specified in \u00a717.",
          "Wire refreshInsights() write-through caching and verify Main\u2192IO\u2192Main boundaries at runtime.",
          "Validate graceful offline fallback against the error contract."]:
    bullet(doc, b)

callout(doc, "Summary: GP4 deliverables are complete and merged into the cumulative SRS. The "
        "architecture is consolidated and ready for API implementation in the next phase with minimal "
        "expected refactor cost.")

out = os.path.join(os.path.dirname(__file__), "BTJ-WeeklyStatusReport-060326-brandonGalli.docx")
doc.save(out)
print("SAVED:", out)
