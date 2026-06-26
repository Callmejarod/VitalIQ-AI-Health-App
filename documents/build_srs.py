# -*- coding: utf-8 -*-
"""Builds the cumulative VitalIQ SRS (GP1-GP3 + Week 4 + Week 5 / GP4) docx."""
import io, sys
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8")
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx_helpers import (new_doc, set_margins, add_table, H1, H2, H3, P, bullet,
                          numbered, code_block, add_image, page_break, callout,
                          add_toc_field, NAVY, BLUE, GREY, TEAL)

doc = new_doc()
set_margins(doc, 1.0)

# ===========================================================================
# COVER
# ===========================================================================
P(doc, "Software Requirements Specification", bold=True, size=30, color=NAVY,
  align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
P(doc, "Cumulative SRS — GP1 + GP2 + GP3 + Week 4 + GP4 (Architectural Consolidation)",
  bold=True, size=13, color=BLUE, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=18)
P(doc, "VitalIQ", bold=True, size=24, color=TEAL, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)
P(doc, "AI-Powered Personal Health Intelligence Platform (Native Android)", size=13,
  align=WD_ALIGN_PARAGRAPH.CENTER, color=GREY, space_after=24)
P(doc, "Team BTJ", bold=True, size=13, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)
P(doc, "Brandon Galli  ·  Jarod Atienzo  ·  Thaigo Amaro Da Silva", size=12,
  align=WD_ALIGN_PARAGRAPH.CENTER, space_after=18)
P(doc, "GitHub Repository:", bold=True, size=11, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)
P(doc, "https://github.com/Callmejarod/VitalIQ-AI-Health-App", size=11,
  color=BLUE, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=18)
P(doc, "Course: DEV322  ·  Milestone GP4 — Architectural Consolidation & Second Data Source Planning",
  size=10.5, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, color=GREY, space_after=2)
P(doc, "Document Version 4.0  ·  Cumulative", size=10.5, italic=True,
  align=WD_ALIGN_PARAGRAPH.CENTER, color=GREY)
page_break(doc)

# ===========================================================================
# DOCUMENT CONTROL / REVISION
# ===========================================================================
H1(doc, "Document Control & Revision History")
P(doc, "This Software Requirements Specification is a cumulative, living document. Each milestone "
  "adds its sections on top of all previously delivered content. Sections from prior milestones "
  "have been retained for continuity and corrected in accordance with instructor feedback "
  "(see \u201cInstructor Feedback Resolution Log\u201d below).")
add_table(doc,
    ["Version", "Milestone", "Date", "Primary Additions"],
    [["1.0", "GP1 — Project Pitch / Scope", "May 6, 2026", "Functional narrative, scope & scale, assumptions, sensor & API justification."],
     ["2.0", "GP2 — Requirements & Stories", "May 13, 2026", "User requirements, user stories, acceptance criteria, use cases."],
     ["3.0", "GP3 — System & Core Architecture", "May 13, 2026", "MVVM layering, ViewModel ownership, coroutine placement, persistence intent, activity/sequence/state diagrams."],
     ["3.5", "Week 4 — Documentation Continuity", "May 20, 2026", "Operational flows, traceability matrix, DAO/ViewModel ownership tables (carried forward)."],
     ["4.0", "GP4 — Architectural Consolidation", "Jun 3, 2026", "Repository Evolution Plan, API Contract, DTO/Mapping Plan, Thread Boundary Plan, dual-source diagrams, Architectural Risk Statement."]],
    widths=[0.8, 2.6, 1.3, 4.2])

H2(doc, "Instructor Feedback Resolution Log")
P(doc, "The following corrections were applied to prior-week content based on instructor comments:")
add_table(doc,
    ["#", "Prior-Week Item", "Instructor Comment", "Correction Applied"],
    [["1", "Activity Diagrams (GP3)", "Steps were not broken down to one activity per step; pseudocode was not derivable.",
      "Re-authored the \u201cList of Steps\u201d so each step contains exactly one activity, making the steps directly usable as pseudocode."],
     ["2", "Activity Diagrams (GP3)", "Main vs IO thread boundaries were not clearly indicated.",
      "All activity diagrams now use explicit two-swimlane Main / IO-Default separation with labeled thread-crossing edges."],
     ["3", "Use Case Diagrams (GP2)", "Use external-actor stick figures where a system communicates with an API/sensor.",
      "Added secondary actors (Device Sensors, AI Insights API) to the relevant use-case diagrams."],
     ["4", "User Stories (GP2)", "Stories must strictly follow the \u201cAs a __ I want __ to achieve __\u201d form.",
      "Normalized every user story into the required three-part template."],
     ["5", "Layering language (GP3)", "Be explicit that the ViewModel must not know the data source.",
      "Strengthened Layer Boundary Rules and added an explicit data-source-agnostic guarantee (carried into the Repository Evolution Plan)."],
     ["6", "Persistence intent (GP3)", "Clarify caching vs transient state for the API feature.",
      "Added explicit persisted-vs-transient field tables (see API Contract & DTO Mapping Plan)."]],
    widths=[0.4, 2.0, 3.0, 3.5], font=9)
page_break(doc)

# ===========================================================================
# TABLE OF CONTENTS
# ===========================================================================
H1(doc, "Table of Contents")
P(doc, "The document is organized cumulatively. Part A retains and corrects all prior-milestone "
  "content (GP1\u2013GP3 and Week 4). Part B contains the new GP4 architectural-consolidation "
  "deliverables.", italic=True, color=GREY)
add_toc_field(doc)
P(doc, "", space_after=2)
P(doc, "Document map (high level):", bold=True, size=10.5, color=BLUE)
for line in [
    "Part A — Prior Milestones (Corrected & Carried Forward)",
    "    1. Functional Narrative — Project Scope & Scale",
    "    2. Assumptions (Platform, Environment, Compliance)",
    "    3. User Requirements (Functional, Non-Functional, System Design)",
    "    4. User Stories",
    "    5. Acceptance Criteria",
    "    6. Use Cases",
    "    7. Use Case Diagrams",
    "    8. Core Architecture — ViewModel Ownership, Coroutine Placement, Persistence Intent",
    "    9. Operational Flows & Step Lists (Pseudocode)",
    "   10. Activity Diagrams (Main vs IO)",
    "   11. Sequence Diagrams",
    "   12. State Machine Diagram",
    "   13. Traceability Matrix",
    "Part B — GP4 Architectural Consolidation (This Milestone)",
    "   14. Coroutine Foundations (Suspend, Main Thread, Scope Ownership)",
    "   15. Repository Evolution Plan",
    "   16. API Contract Specification",
    "   17. DTO + Mapping Plan",
    "   18. Thread Boundary Plan",
    "   19. Updated Activity Diagram (Dual Data Source)",
    "   20. Updated Sequence Diagram (Dual Data Source)",
    "   21. Architectural Risk Statement",
]:
    pp = doc.add_paragraph()
    pp.paragraph_format.space_after = 0
    r = pp.add_run(line)
    r.font.size = __import__("docx").shared.Pt(10)
    if line.startswith("Part"):
        r.bold = True; r.font.color.rgb = NAVY
page_break(doc)

print("cover + control + toc done")

# ===========================================================================
# PART A HEADER
# ===========================================================================
P(doc, "PART A", bold=True, size=22, color=NAVY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)
P(doc, "Prior Milestones — Corrected & Carried Forward (GP1\u2013GP3, Week 4)",
  bold=True, size=13, color=BLUE, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)
P(doc, "Content originally delivered in GP1\u2013GP3 and Week 4, retained for continuity and updated "
  "per the Instructor Feedback Resolution Log.", italic=True, color=GREY,
  align=WD_ALIGN_PARAGRAPH.CENTER)
page_break(doc)

# ---------------------------------------------------------------------------
# 1. FUNCTIONAL NARRATIVE
# ---------------------------------------------------------------------------
H1(doc, "1. Functional Narrative — Project Scope & Scale")
P(doc, "VitalIQ is a native Android application that serves as a comprehensive personal health and "
  "fitness intelligence platform. The system integrates hardware sensor data, manual biometric "
  "logging, and AI-powered analysis to give users a holistic, data-driven picture of their health. "
  "VitalIQ addresses the gap between passive fitness tracking and actionable health intelligence: "
  "while most apps merely record raw data, VitalIQ interprets that data and translates it into "
  "personalized, ranked health recommendations. It accomplishes this by combining real-time sensor "
  "acquisition, structured local persistence, and an external AI inference API into a single, "
  "unified user experience.")
P(doc, "The narrative below describes the entire system from which the prototype's User Stories are "
  "derived. The prototype implements a vertical slice of this system: profile setup, sensor-driven "
  "workout tracking, manual health logging, local persistence, and AI insight generation against an "
  "external API.")

H2(doc, "1.1 Functional Scope")
H3(doc, "Sensor-Based Activity Tracking")
for b in ["Real-time workout tracking using the device accelerometer and step counter via Android SensorManager.",
          "Automatic activity classification into four states: walking, running, stationary, and general movement.",
          "Live display of step count, session duration, movement intensity, and estimated calorie burn during active sessions."]:
    bullet(doc, b)
H3(doc, "Manual Biometric & Lifestyle Logging")
for b in ["User input for blood pressure, resting heart rate, body weight, body composition, and sleep duration.",
          "Medication adherence logging and nutritional intake tracking including water consumption.",
          "Health record import in JSON and CSV formats to consolidate pre-existing user data into the local store."]:
    bullet(doc, b)
H3(doc, "AI-Powered Health Insights")
for b in ["Aggregation of all stored sensor data, manual entries, and historical trends into a structured payload.",
          "Transmission of that payload to an external AI model API (external network API integration).",
          "Display of a holistic health score, category-level breakdowns (cardiovascular, activity, nutrition, hydration, medication adherence, recovery), and ranked improvement suggestions."]:
    bullet(doc, b)
H3(doc, "Data Visualization, History & Profile")
for b in ["Graphical visualization of tracked metrics over time: step trends, weight trend, blood pressure trend, activity breakdown.",
          "Full workout session history including activity type, duration, intensity, and calories burned.",
          "User profile capturing age, height, weight, and fitness goals, used to personalize calorie estimation and AI insight generation."]:
    bullet(doc, b)

H2(doc, "1.2 Out-of-Scope Items")
P(doc, "The following features are explicitly excluded from the current project scope to maintain MVP focus:")
add_table(doc, ["Excluded Feature", "Rationale"],
    [["Wearable / Smartwatch Integration", "Requires additional Bluetooth APIs and hardware pairing logic beyond the defined sensor scope."],
     ["Social Features / Messaging", "Introduces network-layer complexity, authentication, and privacy concerns unrelated to health tracking."],
     ["Cloud Sync / Remote Storage", "All persistence is local via Room. Cloud sync is a future enhancement."],
     ["Medical Diagnosis / Clinical Advice", "VitalIQ provides wellness insights only. No diagnostic or clinical claims are made."]],
    widths=[2.6, 5.0])

# ---------------------------------------------------------------------------
# 2. ASSUMPTIONS
# ---------------------------------------------------------------------------
H1(doc, "2. Assumptions (Platform, Environment, Compliance)")
H2(doc, "2.1 Target Platform")
add_table(doc, ["Parameter", "Assumption"],
    [["Operating System", "Android (native)."],
     ["Minimum SDK", "Android 8.0 (API 26) — ensures availability of required sensor APIs, Room, and coroutines."],
     ["Target SDK", "Android 14 (API 34) — current stable release for full lifecycle and permission compatibility."],
     ["Screen Orientation", "Portrait primary. Landscape not explicitly supported in MVP scope."],
     ["Form Factor", "Smartphone. Tablets and foldables not targeted in MVP scope."]],
    widths=[2.0, 5.6])
H2(doc, "2.2 Hardware Assumptions")
add_table(doc, ["Component", "Assumption"],
    [["Accelerometer", "Present on all target devices and emulators. Required for activity classification; unavailability handled with a user-facing error state."],
     ["Step Counter", "Available on physical devices (API 19+). May be absent on emulators; emulator testing falls back to accelerometer-derived step estimation."],
     ["Internet Connectivity", "Required for AI Insights API calls. App remains fully functional offline with a graceful degraded state."],
     ["Local Storage", "Minimum 50 MB available for the Room database and imported health records."]],
    widths=[2.0, 5.6])
H2(doc, "2.3 Development Environment")
add_table(doc, ["Component", "Specification"],
    [["IDE", "Android Studio"], ["Language", "Kotlin"],
     ["UI Framework", "Jetpack Compose (declarative UI, no XML layouts)"],
     ["Async Framework", "Kotlin Coroutines + Flow / StateFlow"],
     ["Local Database", "Room Persistence Library (SQLite abstraction)"],
     ["Networking (planned)", "Retrofit + Gson (introduced in the next milestone; design only this week)"],
     ["Navigation", "Jetpack Navigation Component with a structured NavGraph"],
     ["Version Control", "Git via GitHub; feature branches with pull requests"],
     ["Project Management", "GitHub Project Board"],
     ["Build System", "Gradle with Kotlin DSL (build.gradle.kts)"]],
    widths=[2.0, 5.6])
H2(doc, "2.4 Compliance & Privacy Assumptions")
for b in ["VitalIQ does not transmit personally identifiable information (PII) to any external server except as included in the AI insights payload (aggregated health metrics only \u2014 no name, email, or location).",
          "The application does not provide medical diagnoses or clinical advice. All AI-generated content is framed as wellness guidance.",
          "Health data is stored exclusively on the user's local device via Room. No cloud backup or remote sync is implemented in this iteration.",
          "The application complies with Android runtime permission requirements (ACTIVITY_RECOGNITION for the step counter on Android 10+)."]:
    bullet(doc, b)
H2(doc, "2.5 Testing Assumptions")
for b in ["Primary testing on the Android Emulator (API 34) via AVD Manager.",
          "Physical-device testing on team members' personal Android devices where available.",
          "Sensor simulation in the emulator for accelerometer events; emulator step-counter limitations are known and accepted.",
          "Manual testing with documented test cases is sufficient for the MVP; no automated UI testing is required this iteration."]:
    bullet(doc, b)
page_break(doc)

print("part A: narrative + assumptions done")

# ---------------------------------------------------------------------------
# 3. USER REQUIREMENTS
# ---------------------------------------------------------------------------
H1(doc, "3. User Requirements")
P(doc, "Requirements are grouped as Functional (FR), Non-Functional (NFR), and System Design (SD). "
  "Functional requirements define what the system does; non-functional requirements define quality "
  "attributes; system-design requirements constrain the architecture.")

H2(doc, "3.1 Functional Requirements")
fr_groups = [
    ("User Account & Profile", [
        ("FR-1", "The system shall allow users to create and edit a personal health profile (name, age, height, weight, fitness goals)."),
        ("FR-2", "The system shall persist user profile information locally using Room."),
        ("FR-3", "The system shall allow users to import existing health records in JSON or CSV format.")]),
    ("Sensor-Based Workout Tracking", [
        ("FR-4", "The system shall access the device accelerometer using Android SensorManager."),
        ("FR-5", "The system shall access the device step counter using Android SensorManager."),
        ("FR-6", "The system shall continuously monitor sensor data during active workout sessions."),
        ("FR-7", "The system shall classify movement into walking, running, stationary, and general movement."),
        ("FR-8", "The system shall display live step counts during workout sessions."),
        ("FR-9", "The system shall track workout session duration in real time."),
        ("FR-10", "The system shall allow users to start, pause, and end a workout session."),
        ("FR-11", "The system shall apply filtering/threshold logic to reduce inaccurate sensor readings.")]),
    ("Workout & Activity Logging", [
        ("FR-12", "The system shall store completed workout sessions locally in Room."),
        ("FR-13", "The system shall display a list of historical workout sessions."),
        ("FR-14", "The system shall persist activity metrics (steps, classification, duration, intensity, calories)."),
        ("FR-15", "The system shall maintain historical data for trend analysis over time.")]),
    ("Health & Lifestyle Logging", [
        ("FR-16", "The system shall allow users to manually log blood pressure, resting heart rate, weight, and body composition."),
        ("FR-17", "The system shall allow users to log medications and adherence."),
        ("FR-18", "The system shall allow users to log food intake and nutrition."),
        ("FR-19", "The system shall allow users to record daily water intake."),
        ("FR-20", "The system shall allow users to record sleep duration and quality.")]),
    ("AI Insights & Intelligence", [
        ("FR-21", "The system shall compile sensor data, workout history, and manual logs into a structured AI payload."),
        ("FR-22", "The system shall send health data to an external AI API for analysis."),
        ("FR-23", "The system shall display an AI-generated overall health score."),
        ("FR-24", "The system shall display AI-generated category scores (cardiovascular, activity, nutrition, hydration, recovery, medication adherence)."),
        ("FR-25", "The system shall display ranked actionable improvement suggestions.")]),
    ("Reactive UI & State", [
        ("FR-26", "The system shall automatically update UI components when sensor or health data changes."),
        ("FR-27", "The system shall expose application state using StateFlow from the ViewModel layer."),
        ("FR-28", "The system shall preserve UI state during configuration changes using ViewModel architecture."),
        ("FR-29", "The system shall display live workout metrics during active tracking sessions.")]),
    ("Navigation & Screens", [
        ("FR-30", "The system shall use Jetpack Navigation for screen transitions."),
        ("FR-31", "The system shall provide a Dashboard screen (daily summaries, health score, quick actions)."),
        ("FR-32", "The system shall provide an Active Workout screen (classification, steps, duration, intensity)."),
        ("FR-33", "The system shall provide a Health Log screen for manual data entry."),
        ("FR-34", "The system shall provide a History & Analytics screen with trend visualization."),
        ("FR-35", "The system shall provide an AI Insights screen."),
        ("FR-36", "The system shall provide a Profile Setup screen with health-record import.")]),
    ("Persistence & Repository", [
        ("FR-37", "The system shall use a Repository layer to separate UI from persistence and API access."),
        ("FR-38", "The system shall persist all core data using Room."),
        ("FR-39", "The system shall allow access to previously stored data without internet connectivity.")]),
    ("Asynchronous Processing", [
        ("FR-40", "The system shall use Kotlin Coroutines for asynchronous execution."),
        ("FR-41", "The system shall perform Room operations on Dispatchers.IO."),
        ("FR-42", "The system shall process sensor computations on Dispatchers.Default."),
        ("FR-43", "The system shall use viewModelScope for lifecycle-safe coroutine execution."),
        ("FR-44", "The system shall ensure sensor, database, and API work never blocks the main thread.")]),
    ("Visualization & Analytics", [
        ("FR-45", "The system shall display historical graphs for steps, weight, blood pressure, and activity trends."),
        ("FR-46", "The system shall display daily aggregate fitness/health summaries."),
        ("FR-47", "The system shall visualize activity distributions across workout classifications.")]),
    ("Constraints & Compliance", [
        ("FR-48", "The system shall maintain data persistence across application restarts."),
        ("FR-49", "The system shall preserve active UI state during device rotation and configuration changes."),
        ("FR-50", "The system shall focus on sensor tracking, health logging, and AI insights, excluding smartwatch integration, social networking, and messaging.")]),
]
fr_rows = []
for grp, items in fr_groups:
    for fid, txt in items:
        fr_rows.append([fid, grp, txt])
add_table(doc, ["ID", "Area", "Requirement"], fr_rows, widths=[0.7, 1.9, 5.0], font=8.5)

H2(doc, "3.2 Non-Functional Requirements")
add_table(doc, ["ID", "Category", "Requirement"],
    [["NFR-1", "Performance", "Sensor processing and UI recomposition shall remain responsive (no perceptible jank) during active tracking."],
     ["NFR-2", "Responsiveness", "No database, network, or sensor operation shall block the main thread; the UI shall remain interactive at all times."],
     ["NFR-3", "Reliability", "Persisted data shall survive process death and configuration changes."],
     ["NFR-4", "Offline Availability", "All core tracking/logging features shall function without connectivity; only AI insight generation requires the network."],
     ["NFR-5", "Privacy", "No PII shall leave the device except aggregated, anonymized metrics in the AI payload."],
     ["NFR-6", "Maintainability", "Swapping the persistence engine or AI provider shall require zero changes to the ViewModel or UI layers."],
     ["NFR-7", "Usability", "Primary actions (start workout, log entry, generate insights) shall be reachable within two taps from the Dashboard."],
     ["NFR-8", "Portability", "The app shall run on Android 8.0+ across mainstream smartphone hardware."]],
    widths=[0.7, 1.5, 5.4], font=9)

H2(doc, "3.3 System Design Requirements")
add_table(doc, ["ID", "Requirement"],
    [["SD-1", "The system shall follow a strict layered MVVM architecture: UI \u2192 ViewModel \u2192 Repository \u2192 (DAO/Room | API). No layer may bypass another."],
     ["SD-2", "The ViewModel shall be the single source of truth for UI state and shall own all coroutine scopes (viewModelScope)."],
     ["SD-3", "The Repository shall be the sole authority over persistence and data-source decisions; it shall hide whether data originates from Room or the API."],
     ["SD-4", "The UI layer shall hold no reference to any Repository, DAO, or database instance and shall launch no data coroutines."],
     ["SD-5", "Network DTOs, domain models, and Room entities shall be separate types connected by explicit mappers."],
     ["SD-6", "Reactive state shall flow unidirectionally: source \u2192 Repository \u2192 ViewModel (StateFlow) \u2192 UI recomposition."]],
    widths=[0.7, 6.9], font=9.5)
page_break(doc)

# ---------------------------------------------------------------------------
# 4. USER STORIES
# ---------------------------------------------------------------------------
H1(doc, "4. User Stories")
P(doc, "Each story follows the required form: \u201cAs a ____, I want to ____, to achieve ____.\u201d "
  "Stories are grouped by epic.", italic=True, color=GREY)

stories = [
    ("Epic: Workout Tracking", "US-1", "Start Workout Session",
     "As a fitness app user, I want to start a workout tracking session, to achieve real-time monitoring of my movement and activity.",
     "High", ["SensorManager integration", "ViewModel state management", "Coroutine sensor processing"]),
    ("Epic: Health Logging", "US-2", "Log Blood Pressure",
     "As a health-conscious user, I want to manually record my blood pressure readings, to achieve the ability to monitor cardiovascular trends over time.",
     "High", ["Room entity for biometric logs", "Repository insertion methods", "Form validation logic"]),
    ("Epic: AI Insights", "US-3", "Generate AI Health Insights",
     "As a user, I want AI-generated health recommendations, to achieve improvement of my workouts and lifestyle habits using personalized guidance.",
     "Medium", ["API integration layer", "Repository networking functions", "AI response parser/mapper"]),
    ("Epic: Profile Management", "US-4", "Create User Profile",
     "As a first-time user, I want to set up my personal profile, to achieve personalization of health tracking and AI analysis.",
     "High", ["Room User entity", "Profile UI screen", "ViewModel state handling"]),
]
current_epic = None
for epic, sid, title, story, prio, deps in stories:
    if epic != current_epic:
        H2(doc, epic); current_epic = epic
    H3(doc, f"{sid} \u2014 {title}")
    P(doc, story)
    P(doc, "Priority: ", bold=True, size=10).add_run(prio)
    P(doc, "Dependencies:", bold=True, size=10)
    for d in deps:
        bullet(doc, d)
page_break(doc)

print("part A: requirements + stories done")

# ---------------------------------------------------------------------------
# 5. ACCEPTANCE CRITERIA
# ---------------------------------------------------------------------------
H1(doc, "5. Acceptance Criteria")
ac = {
    "US-1 \u2014 Start Workout Session": [
        "AC-1.1 User can press a \u201cStart Workout\u201d button.",
        "AC-1.2 Workout timer starts immediately after pressing start.",
        "AC-1.3 Accelerometer data begins collecting in real time.",
        "AC-1.4 Step counter updates continuously during the workout.",
        "AC-1.5 Current activity classification is displayed on screen.",
        "AC-1.6 Workout data updates reactively without manual refresh.",
        "AC-1.7 Workout session persists during screen rotation.",
        "AC-1.8 App handles unavailable sensors gracefully."],
    "US-2 \u2014 Log Blood Pressure": [
        "AC-2.1 User can enter systolic blood pressure value.",
        "AC-2.2 User can enter diastolic blood pressure value.",
        "AC-2.3 User can save the reading successfully.",
        "AC-2.4 Each entry stores the current date and timestamp.",
        "AC-2.5 Data is saved into the Room Database via the Repository.",
        "AC-2.6 Historical BP records display in the history screen.",
        "AC-2.7 Empty or invalid inputs are rejected.",
        "AC-2.8 Validation messages appear for incorrect values."],
    "US-3 \u2014 Generate AI Health Insights": [
        "AC-3.1 App gathers workout and health-log data.",
        "AC-3.2 Structured payload is sent to the AI API.",
        "AC-3.3 AI response includes an overall health score.",
        "AC-3.4 AI response includes category scores.",
        "AC-3.5 AI response includes ranked recommendations.",
        "AC-3.6 AI insights display correctly on screen.",
        "AC-3.7 Loading state appears while waiting for the API response.",
        "AC-3.8 App handles API failures gracefully (falls back to cached/last result).",
        "AC-3.9 Error messages display when requests fail."],
    "US-4 \u2014 Create User Profile": [
        "AC-4.1 User can enter age.",
        "AC-4.2 User can enter height and weight.",
        "AC-4.3 User can select fitness goals.",
        "AC-4.4 Required fields are validated before saving.",
        "AC-4.5 Profile data persists locally after app restart.",
        "AC-4.6 Existing profile data can be edited later.",
        "AC-4.7 Invalid or empty fields display validation errors.",
        "AC-4.8 Saved profile loads automatically when reopening the app."],
}
for story, items in ac.items():
    H2(doc, story)
    for it in items:
        bullet(doc, it)
page_break(doc)

# ---------------------------------------------------------------------------
# 6. USE CASES
# ---------------------------------------------------------------------------
H1(doc, "6. Use Cases")

def use_case(name, actor, goal, pre, trigger, main, alt, post):
    H3(doc, name)
    add_table(doc, ["Field", "Description"],
        [["Use Case Name", name.split("\u2014")[-1].strip()],
         ["Primary Actor", actor],
         ["Goal", goal],
         ["Preconditions", pre],
         ["Trigger", trigger]],
        widths=[1.6, 6.0], font=9.5, zebra=None)
    P(doc, "Main Success Scenario", bold=True, size=10.5, color=BLUE)
    for s in main:
        numbered(doc, s)
    P(doc, "Alternate Flows", bold=True, size=10.5, color=BLUE)
    for a in alt:
        bullet(doc, a)
    P(doc, "Postconditions", bold=True, size=10.5, color=BLUE)
    for p in post:
        bullet(doc, p)

H2(doc, "Workout Tracking Use Cases")
use_case("Use Case 1 \u2014 Start Workout Session", "Fitness App User (secondary: Device Sensors)",
    "Begin tracking workout activity in real time using device sensors.",
    "User has opened the app; sensor permissions granted; sensors available.",
    "User presses the \u201cStart Workout\u201d button.",
    ["User taps \u201cStart Workout.\u201d", "App initializes workout session state.",
     "App registers accelerometer and step counter listeners.", "Workout timer begins running.",
     "App processes sensor data continuously on Dispatchers.Default.",
     "Current activity classification is displayed.", "Live step count updates on screen.",
     "Workout state persists during screen rotation."],
    ["If sensor permissions are denied: app displays a permission request/error.",
     "If sensor hardware is unavailable: app informs the user that tracking cannot start."],
    ["Workout session is active.", "Sensor data is actively collected.", "UI reflects live workout metrics."])

H2(doc, "Health Logging Use Cases")
use_case("Use Case 2 \u2014 Log Blood Pressure Reading", "Health-Conscious User",
    "Record blood pressure readings for historical tracking and analysis.",
    "User is on the blood pressure logging screen; database initialized.",
    "User submits blood pressure values.",
    ["User enters systolic value.", "User enters diastolic value.", "User taps \u201cSave.\u201d",
     "App validates input values.", "App attaches the current timestamp.",
     "Repository stores the entry in Room (Dispatchers.IO).",
     "App refreshes the history/analytics screen.", "New record appears in history."],
    ["If input fields are empty: app displays validation errors.",
     "If values are invalid: app rejects submission and requests correction."],
    ["Blood pressure reading is stored successfully.", "Historical records are updated."])

H2(doc, "AI Insights Use Cases")
use_case("Use Case 3 \u2014 Generate AI Health Insights", "Fitness App User (secondary: AI Insights API)",
    "Receive personalized AI recommendations based on health and workout data.",
    "User has existing workout/health data; internet connection available.",
    "User opens or refreshes the AI Insights screen.",
    ["App gathers workout history and biometric logs from Room.",
     "Repository formats data into a structured API payload.",
     "Repository sends the request to the AI service (Dispatchers.IO).",
     "AI service processes the data.",
     "API returns overall health score, category scores, and ranked recommendations.",
     "Repository maps the response and caches it to Room.",
     "Insights are displayed on the AI Insights screen."],
    ["If the API request fails: app displays an error and falls back to the cached insight.",
     "If response parsing fails: app displays a fallback error state."],
    ["AI recommendations are visible.", "Insights data is cached locally for offline review."])

H2(doc, "Profile Management Use Cases")
use_case("Use Case 4 \u2014 Create User Profile", "First-Time User",
    "Create a personalized profile for health tracking and AI recommendations.",
    "User opens the profile setup screen.",
    "User selects create/setup profile.",
    ["User enters age.", "User enters height.", "User enters weight.", "User selects fitness goals.",
     "User taps \u201cSave Profile.\u201d", "App validates required fields.",
     "Repository stores the profile in Room (Dispatchers.IO).",
     "Profile becomes available throughout the app."],
    ["If required fields are missing: app displays validation messages.",
     "If invalid data is entered: app requests corrected input."],
    ["User profile is stored locally.", "Profile can be edited later."])
page_break(doc)

# ---------------------------------------------------------------------------
# 7. USE CASE DIAGRAMS
# ---------------------------------------------------------------------------
H1(doc, "7. Use Case Diagrams")
P(doc, "Diagrams use the standard stick-figure actor + oval use-case notation. Secondary system "
  "actors (Device Sensors, AI Insights API) are shown where the system communicates with an "
  "external participant.", italic=True, color=GREY)
H2(doc, "Diagram 1 — Workout Tracking")
add_image(doc, "uc1_workout.png", width=6.0, caption="Use Case Diagram 1: Sensor-driven workout tracking.")
H2(doc, "Diagram 2 — Blood Pressure Logging")
add_image(doc, "uc2_bp.png", width=6.0, caption="Use Case Diagram 2: Manual blood pressure logging.")
H2(doc, "Diagram 3 — AI Health Insights")
add_image(doc, "uc3_ai.png", width=6.0, caption="Use Case Diagram 3: AI insight generation via external API.")
H2(doc, "Diagram 4 — User Profile Management")
add_image(doc, "uc4_profile.png", width=6.0, caption="Use Case Diagram 4: User profile creation and editing.")
page_break(doc)

print("part A: acceptance + use cases + uc diagrams done")

# ---------------------------------------------------------------------------
# 8. CORE ARCHITECTURE
# ---------------------------------------------------------------------------
H1(doc, "8. Core Architecture — Ownership, Coroutines & Persistence")
P(doc, "VitalIQ follows a strict layered MVVM architecture. Each layer depends only on the layer "
  "directly below it; no layer bypasses another. This is the architectural spine that GP4 "
  "consolidates before introducing API consumption.")
H2(doc, "8.1 Layered Architecture Overview")
add_table(doc, ["Layer", "Technology", "Responsibility"],
    [["UI", "Jetpack Compose", "Observes StateFlow from the ViewModel. No references to Repository, DAO, or database. Launches no coroutines."],
     ["ViewModel", "viewModelScope + StateFlow", "Owns UI state and lifecycle scope. Launches all coroutines. Calls the Repository only \u2014 never a DAO."],
     ["Repository", "Plain Kotlin classes", "Sole authority over persistence and data-source decisions. Arbitrates between Room and the AI API. Constructs the AI payload."],
     ["DAO", "Room @Dao interfaces", "Exposes suspend functions and Flow queries. Implementations generated by Room at compile time."],
     ["Room Database", "VitalIQDatabase (RoomDatabase)", "Mediates all SQLite access; verifies queries at compile time."],
     ["SQLite", "On-device storage engine", "Physical persistence layer, fully managed by Room, invisible above the DAO."]],
    widths=[1.3, 1.8, 4.5], font=9)

H2(doc, "8.2 ViewModel Ownership")
P(doc, "Each primary screen domain owns a dedicated ViewModel that is the single source of truth "
  "for its UI state. ViewModels interact exclusively with the Repository's public API; they have "
  "no knowledge of any DAO, Room type, or whether data originates from local storage or the API.")
add_table(doc, ["ViewModel", "Owned State", "Primary Trigger"],
    [["WorkoutViewModel", "Active session state, live step count, intensity, duration, classification", "Sensor event callbacks via SensorManager"],
     ["HealthLogViewModel", "Manual entry forms, submission state, entry history", "User form-input events"],
     ["InsightsViewModel", "AI health score, category scores, suggestions, loading/error state", "Repository AI API response"],
     ["HistoryViewModel", "Workout history list, chart series, date-range filter", "Repository Flow emissions from Room"],
     ["ProfileViewModel", "Profile fields, import status, validation errors", "Profile creation/update events"]],
    widths=[1.7, 3.7, 2.2], font=9)

H2(doc, "8.3 Coroutine Placement")
P(doc, "All asynchronous work uses structured concurrency scoped to viewModelScope:")
bullet(doc, "sensor computation, activity classification, calorie estimation.", bold_lead="Dispatchers.Default \u2014 ")
bullet(doc, "all Room reads/writes, AI API requests, and file-import parsing.", bold_lead="Dispatchers.IO \u2014 ")
bullet(doc, "No GlobalScope, no runBlocking, no main-thread blocking.")
bullet(doc, "Coroutine scope ownership belongs exclusively to the ViewModel; the UI launches no data coroutines.")

H2(doc, "8.4 Persistence Intent")
P(doc, "All persistence is handled locally via the Room database (VitalIQDatabase) through the "
  "Repository. No Activity, Composable, or ViewModel touches a DAO directly. The Repository is the "
  "sole owner of persistence decisions. Write operations are suspend functions; queries return Flow "
  "for reactive UI updates.")
add_table(doc, ["Entity", "DAO", "Data Source", "Key Persisted Fields", "Flow"],
    [["User", "UserDao", "Profile Setup screen", "Age, height, weight, fitness goal", "Yes"],
     ["WorkoutSession", "WorkoutSessionDao", "Sensor streams + ViewModel aggregation", "Activity, duration, intensity, steps, calories, timestamp", "Yes"],
     ["ActivityLog", "ActivityLogDao", "Sensor-derived events + import", "Classification, step snapshot, session ID, timestamp", "Yes"],
     ["CachedInsight (GP4)", "InsightDao", "AI API response (write-through cache)", "Health score, category scores, suggestions, timestamp", "Yes"]],
    widths=[1.3, 1.4, 1.9, 2.4, 0.5], font=8.5)
P(doc, "DAO operations summary:", bold=True, size=10.5, color=BLUE)
add_table(doc, ["DAO", "Operation", "Return Type"],
    [["UserDao", "insertUser(user)", "Long (suspend)"],
     ["UserDao", "getUser()", "Flow<User?>"],
     ["WorkoutSessionDao", "insertWorkoutSession(session)", "Long (suspend)"],
     ["WorkoutSessionDao", "deleteWorkoutSession(session)", "Unit (suspend)"],
     ["WorkoutSessionDao", "getAllSessions()", "Flow<List<WorkoutSession>>"],
     ["WorkoutSessionDao", "getSessionById(id)", "WorkoutSession? (suspend)"],
     ["ActivityLogDao", "insertActivityLog(log)", "Long (suspend)"],
     ["ActivityLogDao", "getLogsForSession(sessionId)", "Flow<List<ActivityLog>>"],
     ["ActivityLogDao", "getLogsByDateRange(start, end)", "List<ActivityLog> (suspend)"]],
    widths=[2.0, 3.4, 2.2], font=9)

H2(doc, "8.5 Layer Boundary Rules (Non-Negotiable)")
for b in ["The UI layer holds no reference to any Repository, DAO, or VitalIQDatabase instance.",
          "The ViewModel has no knowledge of Room, SQLite, DAO interfaces, or Room annotations.",
          "The ViewModel does not know whether data originates from Room or the AI Insights API.",
          "The Repository constructs the AI analysis payload (aggregating WorkoutSession, ActivityLog, and User data) before any external call. This logic does not belong in the ViewModel.",
          "Replacing Room with another persistence technology, or swapping the AI provider, requires zero changes to the ViewModel layer."]:
    bullet(doc, b)
page_break(doc)

print("part A: architecture done")

# ---------------------------------------------------------------------------
# 9. OPERATIONAL FLOWS / STEP LISTS (PSEUDOCODE)
# ---------------------------------------------------------------------------
H1(doc, "9. Operational Flows & Step Lists (Pseudocode)")
P(doc, "Per instructor feedback, each flow is decomposed so that every step contains exactly one "
  "activity. When read top-to-bottom, the steps form directly translatable pseudocode for the "
  "annotated activity diagrams that follow. Thread context is annotated as [MAIN], [DEFAULT], or "
  "[IO].", italic=True, color=GREY)

flows = [
    ("9.1 User Profile Creation Flow", [
        "[MAIN] User launches the application.",
        "[IO] Repository queries Room for an existing User profile.",
        "[MAIN] If a profile exists, navigate to the Dashboard.",
        "[MAIN] If no profile exists, display the Profile Setup screen.",
        "[MAIN] User enters age.",
        "[MAIN] User enters height.",
        "[MAIN] User enters weight.",
        "[MAIN] User selects a fitness goal.",
        "[MAIN] User taps Save.",
        "[MAIN] ViewModel validates the input fields.",
        "[MAIN] If validation fails, display an error and remain on the form.",
        "[IO] Repository writes the User entity to Room.",
        "[MAIN] ViewModel updates StateFlow with the new user state.",
        "[MAIN] UI navigates to the Dashboard."]),
    ("9.2 Active Workout Tracking Flow (Sensor-Driven)", [
        "[MAIN] User taps Start Workout on the Dashboard.",
        "[MAIN] ViewModel initializes a WorkoutSession state object.",
        "[MAIN] ViewModel registers the accelerometer listener.",
        "[MAIN] ViewModel registers the step counter listener.",
        "[MAIN] Sensor callback delivers a raw reading.",
        "[DEFAULT] ViewModel filters the movement vector for noise.",
        "[DEFAULT] ViewModel classifies the activity (walking/running/stationary).",
        "[DEFAULT] ViewModel updates the cumulative step count.",
        "[DEFAULT] ViewModel recomputes intensity and duration.",
        "[MAIN] ViewModel emits the updated workout state to StateFlow.",
        "[MAIN] Compose recomposes the live dashboard metrics.",
        "[MAIN] Repeat from step 5 until the user ends the session."]),
    ("9.3 Workout Session Completion Flow", [
        "[MAIN] User taps End Workout.",
        "[MAIN] ViewModel unregisters the SensorManager listeners.",
        "[MAIN] ViewModel finalizes total duration.",
        "[MAIN] ViewModel finalizes total steps.",
        "[MAIN] ViewModel finalizes the activity breakdown.",
        "[MAIN] ViewModel computes estimated calories burned.",
        "[MAIN] ViewModel builds a WorkoutSession domain object.",
        "[IO] Repository persists the WorkoutSession to Room.",
        "[MAIN] StateFlow updates the workout history list.",
        "[MAIN] UI navigates to the History/Summary screen."]),
    ("9.4 Manual Health Data Logging Flow", [
        "[MAIN] User navigates to the Health Log screen.",
        "[MAIN] User enters a biometric/lifestyle value.",
        "[MAIN] User taps Save.",
        "[MAIN] ViewModel validates the data format and required fields.",
        "[MAIN] If invalid, display an error message.",
        "[IO] Repository writes the entry to Room.",
        "[MAIN] ViewModel updates the health-metrics StateFlow.",
        "[MAIN] Compose refreshes the charts and logs."]),
    ("9.5 AI Health Insights Generation Flow", [
        "[MAIN] User opens the AI Insights screen (or taps Generate).",
        "[MAIN] ViewModel sets StateFlow to Loading.",
        "[MAIN] ViewModel calls Repository.refreshInsights() inside viewModelScope.",
        "[IO] Repository reads the cached insight and aggregates logs from Room.",
        "[IO] Repository decides whether the cache is fresh.",
        "[IO] If stale, Repository sends the payload to the AI API.",
        "[IO] Repository maps the API DTO to a domain model.",
        "[IO] Repository writes the mapped insight to the Room cache.",
        "[MAIN] StateFlow emits the domain result (Success/Error).",
        "[MAIN] Compose recomposes the Insights screen."]),
    ("9.6 App Resume / Data Recovery Flow", [
        "[MAIN] User relaunches the app or returns from the background.",
        "[MAIN] The ViewModel is recreated or restored.",
        "[IO] Repository loads the persisted profile from Room.",
        "[IO] Repository loads the historical logs from Room.",
        "[MAIN] StateFlow is repopulated with the latest persisted state.",
        "[MAIN] Compose reconstructs the previous dashboard state.",
        "[MAIN] If a workout was active before termination, offer recovery or discard."]),
]
for title, steps in flows:
    H2(doc, title)
    for s in steps:
        numbered(doc, s)
page_break(doc)

# ---------------------------------------------------------------------------
# 10. ACTIVITY DIAGRAMS
# ---------------------------------------------------------------------------
H1(doc, "10. Activity Diagrams (Main vs IO)")
P(doc, "Each activity diagram uses a two-swimlane structure separating Main-thread activity from "
  "IO/Default-dispatcher activity. Thread-crossing edges (red) mark explicit suspend/resume "
  "boundaries.", italic=True, color=GREY)
H2(doc, "Activity Diagram 1 — Sensor-Driven Workout Tracking")
add_image(doc, "act_workout.png", width=5.7, caption="Workout tracking: sensor callbacks on Main, classification on Default, persistence on IO.")
H2(doc, "Activity Diagram 2 — AI Insights Generation (preview of GP4 dual-source)")
add_image(doc, "act_insights_week5.png", width=5.2, caption="AI insights with Repository branching and Main\u2192IO\u2192Main boundaries (expanded in \u00a719).")
page_break(doc)

# ---------------------------------------------------------------------------
# 11. SEQUENCE DIAGRAMS
# ---------------------------------------------------------------------------
H1(doc, "11. Sequence Diagrams")
P(doc, "The sequence diagram below models how work moves through the system, where suspend "
  "boundaries occur, which layer owns each responsibility, and how state returns to the UI "
  "reactively. It is expanded in \u00a720 for the dual-data-source architecture.", italic=True, color=GREY)
add_image(doc, "seq_insights_week5.png", width=6.6, caption="Sequence: AI insights across UI \u2192 ViewModel \u2192 Repository \u2192 (Room | API) with thread bands.")
P(doc, "Coroutine Ownership: ", bold=True, size=10).add_run(
    "all coroutines launch in viewModelScope (InsightsViewModel); navigation away cancels the call.")
P(doc, "Suspend Boundaries: ", bold=True, size=10).add_run(
    "the entire repository chain runs under withContext(Dispatchers.IO); the UI thread is never blocked.")
P(doc, "Response Mapping: ", bold=True, size=10).add_run(
    "raw JSON DTO is mapped to a domain model inside the Repository, on the IO dispatcher.")
P(doc, "Flow Emission: ", bold=True, size=10).add_run(
    "the result flows back through the ViewModel StateFlow, triggering Compose recomposition on Main.")
page_break(doc)

# ---------------------------------------------------------------------------
# 12. STATE MACHINE DIAGRAM
# ---------------------------------------------------------------------------
H1(doc, "12. State Machine Diagram")
P(doc, "This state machine represents the behavioral core of the system because workout tracking is "
  "the primary real-time feature that drives sensor processing, reactive UI updates, persistence, "
  "and AI analysis. The system transitions between Idle, Active, Paused, Saving, and Completed based "
  "on user actions and sensor events. These transitions control how data flows through the "
  "ViewModel, StateFlow, and Room persistence layers.")
add_image(doc, "state_workout.png", width=6.2, caption="Workout session lifecycle; the Saving\u2192Completed transition executes on Dispatchers.IO.")
page_break(doc)

# ---------------------------------------------------------------------------
# 13. TRACEABILITY MATRIX
# ---------------------------------------------------------------------------
H1(doc, "13. Traceability Matrix")
P(doc, "Maps user stories to acceptance criteria to ensure complete coverage.")
trace = [
    ("US-1", "Start Workout Session", "AC-1.1", "User can press a Start Workout button."),
    ("US-1", "Start Workout Session", "AC-1.2", "Workout timer starts immediately."),
    ("US-1", "Start Workout Session", "AC-1.3", "Accelerometer data begins collecting in real time."),
    ("US-1", "Start Workout Session", "AC-1.4", "Step counter updates continuously."),
    ("US-1", "Start Workout Session", "AC-1.5", "Current activity classification is displayed."),
    ("US-1", "Start Workout Session", "AC-1.7", "Workout state persists during screen rotation."),
    ("US-2", "Log Blood Pressure", "AC-2.1", "User can enter systolic value."),
    ("US-2", "Log Blood Pressure", "AC-2.2", "User can enter diastolic value."),
    ("US-2", "Log Blood Pressure", "AC-2.4", "Entry includes timestamp/date."),
    ("US-2", "Log Blood Pressure", "AC-2.5", "Data is stored in Room via the Repository."),
    ("US-2", "Log Blood Pressure", "AC-2.6", "Historical records display in history."),
    ("US-2", "Log Blood Pressure", "AC-2.7", "Invalid inputs are rejected."),
    ("US-3", "Generate AI Insights", "AC-3.2", "Structured payload is sent to the AI API."),
    ("US-3", "Generate AI Insights", "AC-3.3", "Response includes overall health score."),
    ("US-3", "Generate AI Insights", "AC-3.4", "Response includes category scores."),
    ("US-3", "Generate AI Insights", "AC-3.5", "Response includes ranked recommendations."),
    ("US-3", "Generate AI Insights", "AC-3.8", "API failures handled gracefully via cache fallback."),
    ("US-4", "Create User Profile", "AC-4.1", "User can enter age."),
    ("US-4", "Create User Profile", "AC-4.2", "User can enter height and weight."),
    ("US-4", "Create User Profile", "AC-4.3", "User can select fitness goals."),
    ("US-4", "Create User Profile", "AC-4.4", "Required fields are validated."),
    ("US-4", "Create User Profile", "AC-4.5", "Profile data persists locally."),
    ("US-4", "Create User Profile", "AC-4.6", "Existing profile can be edited later."),
]
add_table(doc, ["Story ID", "User Story", "AC ID", "Acceptance Criteria"],
          [list(r) for r in trace], widths=[0.8, 2.0, 0.8, 4.0], font=8.5)
page_break(doc)

print("part A complete")

# ===========================================================================
# PART B HEADER
# ===========================================================================
P(doc, "PART B", bold=True, size=22, color=NAVY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)
P(doc, "GP4 — Architectural Consolidation & Second Data Source Planning",
  bold=True, size=13, color=BLUE, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)
callout(doc, "Scope note: This is a structural / architectural-design checkpoint. No networking or "
        "Retrofit code is required or added in this milestone. We are designing the contract and the "
        "Repository evolution plan that will drive API consumption in the next phase. No navigation "
        "changes and no new frameworks are introduced.")
page_break(doc)

# ---------------------------------------------------------------------------
# 14. COROUTINE FOUNDATIONS
# ---------------------------------------------------------------------------
H1(doc, "14. Coroutine Foundations")
P(doc, "Before describing how the Repository evolves, this section establishes the coroutine "
  "literacy that the rest of Part B depends on. It answers three questions directly.")

H2(doc, "14.1 Why suspend Exists")
P(doc, "A suspend function is a function that can pause its execution and resume later without "
  "blocking the underlying thread. It exists to model long-running, asynchronous work (network "
  "calls, disk I/O, database queries) as ordinary sequential code while freeing the thread during "
  "the wait.")
for b in ["Without suspend, the only ways to wait for I/O are to block a thread (wasting it) or to "
          "fragment logic into nested callbacks. suspend gives us the readability of sequential code "
          "with the efficiency of asynchronous execution.",
          "A suspend function can only be called from another suspend function or a coroutine builder "
          "(launch/async). This compiler-enforced rule guarantees that suspending work is always "
          "anchored to a coroutine scope and a dispatcher.",
          "At a suspension point the coroutine saves its state (a continuation) and releases the "
          "thread back to the dispatcher's pool. When the awaited work completes, the coroutine is "
          "resumed \u2014 possibly on a different thread \u2014 with its state intact.",
          "In VitalIQ, every Room DAO write and the AI API call are suspend functions. This lets the "
          "Repository express \u201cread cache, then call API, then write cache\u201d as straight-line "
          "code while never holding a thread idle during the HTTP round-trip."]:
    bullet(doc, b)

H2(doc, "14.2 Why the Main Thread Cannot Block")
P(doc, "On Android the main (UI) thread is responsible for rendering frames, dispatching input "
  "events, and \u2014 with Jetpack Compose \u2014 running recomposition. It must complete each frame "
  "within roughly 16 ms to sustain 60 fps.")
for b in ["If the main thread performs a blocking operation (a synchronous network call, a disk read, "
          "a heavy loop), it cannot render frames or respond to touch. The UI freezes.",
          "Android enforces this: a network call on the main thread throws NetworkOnMainThreadException, "
          "and prolonged blocking triggers an Application Not Responding (ANR) dialog.",
          "Therefore VitalIQ keeps all Room I/O, AI API calls, and file parsing off the main thread by "
          "running them under Dispatchers.IO, and sensor math under Dispatchers.Default. Only fast "
          "state emission and recomposition happen on Main.",
          "suspend + dispatchers are the mechanism that makes this possible: the ViewModel launches a "
          "coroutine on Main, the work suspends and hops to IO, and only the final result resumes on "
          "Main to update StateFlow."]:
    bullet(doc, b)

H2(doc, "14.3 Why Scope Ownership Matters")
P(doc, "A CoroutineScope defines the lifetime and cancellation boundary of the coroutines launched "
  "within it. Owning the correct scope is what makes concurrency safe and leak-free.")
for b in ["viewModelScope is tied to the ViewModel's lifecycle. When the ViewModel is cleared "
          "(onCleared), every coroutine in that scope is cancelled automatically. This prevents work "
          "from outliving the screen that started it.",
          "Structured concurrency means child coroutines are bound to a parent scope: if the parent is "
          "cancelled, children are cancelled; if a child fails, the failure propagates predictably. "
          "There are no orphaned background tasks.",
          "VitalIQ forbids GlobalScope and runBlocking. GlobalScope creates coroutines that never get "
          "cancelled with the screen (leaks and wasted work, e.g. an AI call continuing after the user "
          "navigates away). runBlocking blocks the calling thread, defeating the purpose of suspend.",
          "Because the ViewModel owns the scope, the Repository and DAO never create their own scopes. "
          "They expose suspend functions and Flows; the ViewModel decides when, where, and for how long "
          "they run. This keeps cancellation, lifecycle safety, and thread policy in exactly one place."]:
    bullet(doc, b)
page_break(doc)

print("part B: coroutine foundations done")

# ---------------------------------------------------------------------------
# 15. REPOSITORY EVOLUTION PLAN
# ---------------------------------------------------------------------------
H1(doc, "15. Repository Evolution Plan")
P(doc, "Today the Repository mediates between the ViewModel and Room. This milestone formalizes how "
  "that Repository will expand to support a second data source (the AI Insights / health-data API) "
  "while preserving MVVM discipline, coroutine boundaries, and reactive state flow. No networking "
  "code is added now; this is the contract and evolution plan that will drive the API phase.")

H2(doc, "15.1 Current Repository Functions (Room-only)")
P(doc, "The Repository currently exposes a Room-backed surface. ViewModels call only these functions:")
code_block(doc,
"interface VitalIQRepository {\n"
"    // Profile\n"
"    fun observeProfile(): Flow<User?>\n"
"    suspend fun saveProfile(user: User)\n\n"
"    // Workouts\n"
"    fun observeWorkouts(): Flow<List<WorkoutSession>>\n"
"    suspend fun saveWorkout(session: WorkoutSession)\n"
"    suspend fun getWorkout(id: Int): WorkoutSession?\n\n"
"    // Activity / Health logs\n"
"    fun observeActivityLogs(): Flow<List<ActivityLog>>\n"
"    suspend fun logActivity(log: ActivityLog)\n"
"    suspend fun logsByDateRange(start: Long, end: Long): List<ActivityLog>\n"
"}")
P(doc, "Internally these delegate to UserDao, WorkoutSessionDao, and ActivityLogDao, all executed "
  "under withContext(Dispatchers.IO).")

H2(doc, "15.2 New Functions Required for API Integration")
P(doc, "API support adds remote-capable functions to the same Repository interface. The function "
  "signatures expose only domain types \u2014 never DTOs or Room entities.")
code_block(doc,
"interface VitalIQRepository {\n"
"    // ... existing Room functions unchanged ...\n\n"
"    // NEW: reactive cached insight (single source of truth for UI)\n"
"    fun observeInsights(): Flow<HealthInsight?>      // emits from Room cache\n\n"
"    // NEW: refresh from remote, write-through to cache, then Flow re-emits\n"
"    suspend fun refreshInsights(): Result<HealthInsight>\n\n"
"    // NEW: remote sync helpers (still domain-typed)\n"
"    suspend fun syncProfileToRemote(): Result<Unit>\n"
"    suspend fun pullRemoteWorkouts(): Result<Unit>\n"
"}")
P(doc, "Key design choice: the UI observes observeInsights() (a Room-backed Flow), while "
  "refreshInsights() is a one-shot command that fetches remotely, maps, and writes to the cache. "
  "The cache write causes the Flow to re-emit. This is the single-source-of-truth pattern: the UI "
  "always renders from Room; the network only updates the cache.")

H2(doc, "15.3 What Must Remain Unchanged in the ViewModel")
for b in ["The ViewModel continues to call only the Repository's public, domain-typed API.",
          "The ViewModel still owns viewModelScope and launches all coroutines there.",
          "The ViewModel still exposes a single StateFlow of domain state to the UI.",
          "No DTO, Retrofit, OkHttp, or Room type ever appears in the ViewModel \u2014 before or after the API is added.",
          "Adding the API requires no new method on the ViewModel beyond what already maps to user intent (e.g. load() / generateInsights())."]:
    bullet(doc, b)

H2(doc, "15.4 What Must Remain Unchanged in the UI")
for b in ["Composables continue to observe StateFlow via collectAsState() and recompose; they launch no coroutines.",
          "The UI holds no reference to the Repository, DAO, ApiService, or database.",
          "The UI renders the same domain/UI models regardless of whether the data came from Room or the API.",
          "No navigation graph changes are introduced by the API."]:
    bullet(doc, b)

H2(doc, "15.5 How the Repository Mediates Between Room and API")
P(doc, "The Repository is the single mediator. It decides, per request, which source to use and how "
  "to reconcile them, exposing only domain types upward.")
add_table(doc, ["Concern", "Repository Responsibility"],
    [["Source selection", "Decide local-only, remote-only, or cache-then-network per function (see branching logic below)."],
     ["Mapping", "Convert DTO \u2192 Domain \u2192 Entity (and reverse) so callers never see network or persistence types."],
     ["Caching", "Write-through remote results to Room so the UI's Flow re-emits and offline reads succeed."],
     ["Threading", "Wrap all I/O in withContext(Dispatchers.IO); never touch the main thread."],
     ["Error handling", "Catch network/parse errors and return a domain Result; fall back to the cached value."],
     ["Single source of truth", "Expose Room-backed Flows for observed data; the network mutates the cache, not the UI directly."]],
    widths=[1.8, 5.8], font=9.5)

H2(doc, "15.6 Where Branching Logic Lives")
P(doc, "All data-source branching lives inside the Repository \u2014 never in the ViewModel, UI, or "
  "DAO. The canonical cache-then-network pattern:")
code_block(doc,
"class VitalIQRepositoryImpl(\n"
"    private val insightDao: InsightDao,        // Room (local)\n"
"    private val api: ApiService,               // API (remote)\n"
"    private val mapper: InsightMapper,\n"
"    private val io: CoroutineDispatcher = Dispatchers.IO,\n"
") : VitalIQRepository {\n\n"
"    override fun observeInsights(): Flow<HealthInsight?> =\n"
"        insightDao.observeLatest().map { it?.toDomain() }   // always from cache\n\n"
"    override suspend fun refreshInsights(): Result<HealthInsight> = withContext(io) {\n"
"        val cached = insightDao.getLatest()\n"
"        // ---- BRANCHING LOGIC LIVES HERE ----\n"
"        if (cached != null && cached.isFresh()) {\n"
"            Result.success(cached.toDomain())               // local path\n"
"        } else try {\n"
"            val dto = api.generateInsights()                 // remote path\n"
"            val domain = mapper.dtoToDomain(dto)\n"
"            insightDao.upsert(mapper.domainToEntity(domain)) // write-through cache\n"
"            Result.success(domain)\n"
"        } catch (e: Exception) {\n"
"            cached?.let { Result.success(it.toDomain()) }    // graceful fallback\n"
"                ?: Result.failure(e)\n"
"        }\n"
"    }\n"
"}")
P(doc, "Because the decision \u201cuse cache vs. call API\u201d is encapsulated here, the ViewModel "
  "cannot know \u2014 and never asks \u2014 which source served the data. This is precisely how we "
  "prevent the ViewModel from knowing which data source is used.")
page_break(doc)

print("part B: repository evolution done")

# ---------------------------------------------------------------------------
# 16. API CONTRACT SPECIFICATION
# ---------------------------------------------------------------------------
H1(doc, "16. API Contract Specification")
P(doc, "This contract drives DTOs, mappers, Room entities, and UI models. It is design only \u2014 "
  "no Retrofit code yet. The base URL is {BACKEND_URL}/api/. All endpoints exchange JSON; the AI "
  "insight endpoints are the primary second data source that influences app behavior.")
P(doc, "Endpoint summary:", bold=True, size=10.5, color=BLUE)
add_table(doc, ["Method", "Path", "Purpose"],
    [["GET", "/api/profile", "Fetch the user profile."],
     ["PUT", "/api/profile", "Create/update the user profile."],
     ["GET", "/api/workouts", "List workout sessions."],
     ["POST", "/api/workouts", "Create a workout session."],
     ["GET", "/api/health-entries", "List biometric entries (optional entry_type filter)."],
     ["POST", "/api/health-entries", "Create a biometric entry."],
     ["GET / POST", "/api/medications, /api/nutrition, /api/sleep", "List/create lifestyle logs."],
     ["GET", "/api/dashboard/summary", "Aggregated daily summary."],
     ["POST", "/api/insights/generate", "Generate fresh AI insights from stored data."],
     ["GET", "/api/insights/latest", "Fetch the most recent AI insight."],
     ["POST", "/api/import", "Multipart upload of a JSON/CSV health record."]],
    widths=[1.1, 3.0, 3.5], font=9)

# --- Detailed endpoint 1: generate insights ---
H2(doc, "16.1 Endpoint — Generate AI Insights")
add_table(doc, ["Attribute", "Specification"],
    [["Endpoint URL", "{BACKEND_URL}/api/insights/generate"],
     ["HTTP Method", "POST"],
     ["Path Parameters", "None"],
     ["Query Parameters", "None"],
     ["Request Body", "None (server aggregates the user's stored data server-side). Future: optional client-aggregated payload."],
     ["Success Response", "200 OK \u2014 InsightDto (see schema below)"],
     ["Error Response", "4xx/5xx with a JSON error envelope (see \u00a716.4)"]],
    widths=[1.7, 5.9], font=9.5, zebra=None)
P(doc, "Response schema (field types):", bold=True, size=10, color=BLUE)
add_table(doc, ["Field", "JSON Type", "Kotlin Type", "Persist?"],
    [["overall_score", "integer", "Int", "Persisted"],
     ["category_scores", "object<string,int>", "Map<String,Int>", "Persisted"],
     ["suggestions", "array<object>", "List<SuggestionDto>", "Persisted"],
     ["suggestions[].title", "string", "String", "Persisted"],
     ["suggestions[].detail", "string", "String", "Persisted"],
     ["suggestions[].priority", "integer", "Int", "Persisted"],
     ["suggestions[].category", "string", "String", "Persisted"],
     ["trend_summary", "string", "String", "Persisted"],
     ["snapshot", "object", "Map<String,Any>", "Transient (debug context)"],
     ["used_fallback", "boolean", "Boolean", "Transient (UI badge only)"],
     ["empty", "boolean", "Boolean", "Transient (controls Empty state)"]],
    widths=[2.0, 1.8, 2.0, 1.8], font=9)
P(doc, "Example JSON response:", bold=True, size=10, color=BLUE)
code_block(doc,
'{\n'
'  "overall_score": 78,\n'
'  "category_scores": {\n'
'    "cardiovascular": 82, "activity": 74, "nutrition": 69,\n'
'    "hydration": 88, "recovery": 71, "medication_adherence": 90\n'
'  },\n'
'  "suggestions": [\n'
'    { "title": "Add a recovery day", "detail": "Your intensity has\n'
'      trended high 5 days running. Schedule active recovery.",\n'
'      "priority": 1, "category": "recovery" },\n'
'    { "title": "Increase water intake", "detail": "You averaged\n'
'      1.4 L/day vs a 2.0 L goal.", "priority": 2, "category": "hydration" }\n'
'  ],\n'
'  "trend_summary": "Activity up 12% week-over-week; hydration below goal.",\n'
'  "snapshot": { "workouts_considered": 9, "logs_considered": 21 },\n'
'  "used_fallback": false,\n'
'  "empty": false\n'
'}')

# --- Detailed endpoint 2: workouts ---
H2(doc, "16.2 Endpoint — Workouts")
add_table(doc, ["Attribute", "Specification"],
    [["Endpoint URL", "{BACKEND_URL}/api/workouts"],
     ["HTTP Methods", "GET (list, optional ?limit=N), POST (create)"],
     ["Path Parameters", "None"],
     ["Query Parameters", "limit: integer (GET only, default 50)"],
     ["Request Body (POST)", "WorkoutDto JSON"],
     ["Success Response", "GET \u2192 array<WorkoutDto>; POST \u2192 created WorkoutDto"],
     ["Error Response", "JSON error envelope (\u00a716.4)"]],
    widths=[1.7, 5.9], font=9.5, zebra=None)
P(doc, "Field types & persistence:", bold=True, size=10, color=BLUE)
add_table(doc, ["Field", "JSON Type", "Persist?"],
    [["id", "string (server)", "Persisted (remote id; local PK is autogen)"],
     ["started_at / ended_at", "string (ISO-8601)", "Persisted"],
     ["duration_seconds", "integer", "Persisted"],
     ["activity_type", "string", "Persisted"],
     ["steps", "integer", "Persisted"],
     ["avg_intensity", "float", "Persisted"],
     ["activity_breakdown", "object<string,int>", "Persisted (serialized JSON column)"]],
    widths=[2.2, 2.2, 3.2], font=9)
P(doc, "Example JSON response:", bold=True, size=10, color=BLUE)
code_block(doc,
'{\n'
'  "id": "wk_01H...",\n'
'  "started_at": "2026-06-03T07:12:00Z",\n'
'  "ended_at":   "2026-06-03T07:46:30Z",\n'
'  "duration_seconds": 2070,\n'
'  "activity_type": "running",\n'
'  "steps": 4120,\n'
'  "avg_intensity": 0.73,\n'
'  "activity_breakdown": { "running": 1500, "walking": 480, "stationary": 90 }\n'
'}')

# --- Detailed endpoint 3: profile ---
H2(doc, "16.3 Endpoint — Profile")
add_table(doc, ["Attribute", "Specification"],
    [["Endpoint URL", "{BACKEND_URL}/api/profile"],
     ["HTTP Methods", "GET (fetch), PUT (upsert)"],
     ["Request Body (PUT)", "ProfileDto JSON"],
     ["Success Response", "ProfileDto JSON"]],
    widths=[1.7, 5.9], font=9.5, zebra=None)
add_table(doc, ["Field", "JSON Type", "Persist?"],
    [["user_id", "string", "Persisted"],
     ["name", "string", "Persisted (local only; not sent in AI payload)"],
     ["age", "integer", "Persisted"],
     ["height_cm / weight_kg", "float", "Persisted"],
     ["fitness_goal", "string", "Persisted"],
     ["daily_step_goal / daily_water_goal_ml", "integer", "Persisted"],
     ["updated_at", "string (ISO-8601)", "Persisted (sync conflict resolution)"]],
    widths=[2.6, 1.8, 3.2], font=9)

H2(doc, "16.4 Error Response Format")
P(doc, "All endpoints return a consistent error envelope so the Repository can map failures to a "
  "single domain error type:")
code_block(doc,
'{\n'
'  "error": {\n'
'    "code": "VALIDATION_ERROR",        // machine-readable\n'
'    "message": "systolic must be 60-260",\n'
'    "status": 422,\n'
'    "field": "systolic"               // optional\n'
'  }\n'
'}')
add_table(doc, ["HTTP Status", "code", "Repository Handling"],
    [["400 / 422", "VALIDATION_ERROR", "Map to DomainError.Validation; surface field message to UI."],
     ["401 / 403", "AUTH_ERROR", "Map to DomainError.Auth; (future auth phase)."],
     ["404", "NOT_FOUND", "Treat as empty; fall back to cache."],
     ["429", "RATE_LIMITED", "Map to DomainError.Retryable; keep cached insight."],
     ["5xx / timeout / offline", "SERVER_ERROR / NETWORK", "Map to DomainError.Network; serve cached insight, show banner."]],
    widths=[1.6, 2.0, 4.0], font=9)
page_break(doc)

print("part B: api contract done")

# ---------------------------------------------------------------------------
# 17. DTO + MAPPING PLAN
# ---------------------------------------------------------------------------
H1(doc, "17. DTO + Mapping Plan")
P(doc, "Three distinct model families exist, one per architectural layer. They are never merged, "
  "because each answers to a different authority: the network, the business rules, and the "
  "database schema.")

H2(doc, "17.1 The Three Models")
add_table(doc, ["Model", "Layer", "Owns / Reflects", "Example (Insight)"],
    [["DTO", "Network", "The wire format / JSON contract. @SerializedName, nullable, snake_case.", "InsightDto(overall_score, category_scores, suggestions, used_fallback, empty)"],
     ["Domain", "Business logic", "Clean, non-nullable Kotlin model the app reasons about. No annotations.", "HealthInsight(score: Int, categories: Map, suggestions: List<Suggestion>, generatedAt: Instant)"],
     ["Entity", "Persistence", "The Room table schema. @Entity, @PrimaryKey, column types.", "CachedInsightEntity(id, scoreJson, suggestionsJson, generatedAtEpoch)"]],
    widths=[0.9, 1.3, 2.7, 2.7], font=8.5)

H2(doc, "17.2 Mapping Path: DTO \u2192 Domain \u2192 Entity")
P(doc, "The forward path is taken when reading from the network and caching:")
code_block(doc,
"// network -> business model\n"
"fun InsightDto.toDomain(): HealthInsight = HealthInsight(\n"
"    score        = overallScore,\n"
"    categories   = categoryScores,\n"
"    suggestions  = suggestions.map { Suggestion(it.title, it.detail, it.priority, it.category) },\n"
"    trend        = trendSummary,\n"
"    generatedAt  = Instant.now(),\n"
")\n\n"
"// business model -> persistence row (serialize collections to JSON columns)\n"
"fun HealthInsight.toEntity(): CachedInsightEntity = CachedInsightEntity(\n"
"    id              = LATEST_ID,\n"
"    overallScore    = score,\n"
"    categoryJson    = gson.toJson(categories),\n"
"    suggestionsJson = gson.toJson(suggestions),\n"
"    trend           = trend,\n"
"    generatedAtEpoch = generatedAt.toEpochMilli(),\n"
")")
P(doc, "Reverse mapping (Entity \u2192 Domain) is required because the UI always reads from the Room "
  "cache (single source of truth). When the cached Flow emits, entities are mapped back to domain "
  "models for the ViewModel:")
code_block(doc,
"fun CachedInsightEntity.toDomain(): HealthInsight = HealthInsight(\n"
"    score       = overallScore,\n"
"    categories  = gson.fromJson(categoryJson, mapType),\n"
"    suggestions = gson.fromJson(suggestionsJson, listType),\n"
"    trend       = trend,\n"
"    generatedAt = Instant.ofEpochMilli(generatedAtEpoch),\n"
")")
P(doc, "A reverse Domain \u2192 DTO path is also defined for write/sync endpoints (e.g. PUT /profile, "
  "POST /workouts), converting domain objects into the JSON the API expects.")

H2(doc, "17.3 Why Mapping Does Not Occur in the UI or DAO Layer")
for b in ["Mapping in the UI would force Composables to import DTO and Entity types, coupling the "
          "view to the network contract and the database schema \u2014 a direct violation of the layer "
          "boundary rules. A backend field rename would then ripple into UI code.",
          "Mapping in the DAO is impossible/improper: the DAO's job is purely to read/write entities. "
          "It has no knowledge of DTOs or business rules, and Room generates its implementation. "
          "Putting network concerns there would bypass the Repository.",
          "The Repository is the only layer that legitimately sees all three models, so it is the "
          "correct and single home for mapping. This keeps each layer ignorant of the others' types "
          "and means a change in the JSON contract is absorbed entirely inside the mapper.",
          "Concentrating mapping in one place also makes it unit-testable in isolation, independent of "
          "Android, Room, or the network."]:
    bullet(doc, b)
page_break(doc)

# ---------------------------------------------------------------------------
# 18. THREAD BOUNDARY PLAN
# ---------------------------------------------------------------------------
H1(doc, "18. Thread Boundary Plan")
P(doc, "This section demonstrates coroutine literacy by stating exactly where each kind of work "
  "executes, on which dispatcher, and why.")
add_table(doc, ["Work", "Where it executes", "Dispatcher", "Why"],
    [["UI trigger / state emission / recomposition", "ViewModel + Compose", "Dispatchers.Main (viewModelScope default)",
      "Touch handling and Compose recomposition must run on the UI thread; emissions are cheap."],
     ["Network IO (AI API, sync)", "Repository (via ApiService suspend fns)", "Dispatchers.IO",
      "HTTP is blocking/latency-bound; IO has an elastic thread pool sized for waiting work. Never on Main (ANR / NetworkOnMainThreadException)."],
     ["Database IO (Room reads/writes)", "Repository (via DAO suspend fns + Flow)", "Dispatchers.IO",
      "Disk access is blocking. Room enforces off-main execution for suspend/Flow queries."],
     ["DTO \u2192 Domain \u2192 Entity mapping", "Repository", "Dispatchers.IO (same withContext block as the IO it serves)",
      "Mapping is co-located with the IO that produced the data, avoiding an extra thread hop; it is light CPU work bounded by payload size."],
     ["Sensor math / activity classification", "WorkoutViewModel processing", "Dispatchers.Default",
      "CPU-bound work belongs on the Default pool (sized to cores), keeping Main free for rendering."]],
    widths=[2.0, 1.9, 1.7, 2.0], font=8.5)
P(doc, "Canonical boundary in code:", bold=True, size=10.5, color=BLUE)
code_block(doc,
"// ViewModel (Main) ---------------------------------------------------\n"
"fun generateInsights() = viewModelScope.launch {        // starts on Main\n"
"    _uiState.value = Loading                            // Main: cheap emit\n"
"    val result = repo.refreshInsights()                 // suspends; no block\n"
"    _uiState.value = result.fold(::Success, ::toError)  // resumes on Main\n"
"}\n\n"
"// Repository (IO) -----------------------------------------------------\n"
"suspend fun refreshInsights(): Result<HealthInsight> =\n"
"    withContext(Dispatchers.IO) {                       // hop Main -> IO\n"
"        val dto    = api.generateInsights()             // IO: network\n"
"        val domain = dto.toDomain()                     // IO: mapping\n"
"        insightDao.upsert(domain.toEntity())            // IO: database\n"
"        Result.success(domain)\n"
"    }                                                   // hop IO -> Main on return")
P(doc, "Boundary guarantees: the coroutine begins and resumes on Main; the suspend call hops to IO "
  "for all network, database, and mapping work; only the lightweight StateFlow assignment happens "
  "back on Main. The main thread is never blocked, and the dispatcher choice is explicit at the "
  "boundary, not scattered across layers.")
page_break(doc)

# ---------------------------------------------------------------------------
# 19. UPDATED ACTIVITY DIAGRAM
# ---------------------------------------------------------------------------
H1(doc, "19. Updated Activity Diagram (Dual Data Source)")
P(doc, "The updated activity diagram shows the UI trigger, the ViewModel call, the Repository "
  "branching logic, both Room and API interactions, the mapping step, the flow emission back to "
  "the UI, and explicit Main vs IO thread boundaries.")
add_image(doc, "act_insights_week5.png", width=5.3, caption="Dual-source AI insights. Red edges cross the Main/IO boundary; the diamond is the Repository's cache-vs-network decision.")
P(doc, "Reading the diagram against the layer rules:", bold=True, size=10.5, color=BLUE)
for b in ["UI trigger and ViewModel launch are in the Main lane; the only Main-thread work is "
          "emitting Loading and, later, the result.",
          "Repository.refreshInsights() crosses into the IO lane via withContext(Dispatchers.IO).",
          "The decision diamond (cache fresh?) is the branching logic \u2014 it lives inside the Repository.",
          "Both the Room path (read/upsert cache) and the API path (suspend HTTP) are in the IO lane.",
          "Mapping DTO\u2192Domain\u2192Entity is shown in the IO lane, adjacent to the IO that produced the data.",
          "Flow emission returns to the Main lane, where Compose recomposes."]:
    bullet(doc, b)
page_break(doc)

# ---------------------------------------------------------------------------
# 20. UPDATED SEQUENCE DIAGRAM
# ---------------------------------------------------------------------------
H1(doc, "20. Updated Sequence Diagram (Dual Data Source)")
P(doc, "The updated sequence diagram reflects the dual-data-source architecture and Repository "
  "mediation. It shows the coroutine scope origin and the thread transitions across the call.")
add_image(doc, "seq_insights_week5.png", width=6.7, caption="Repository mediates between Room (cache) and ApiService (remote); thread bands show Main \u2192 IO \u2192 Main.")
P(doc, "The diagram explicitly shows:", bold=True, size=10.5, color=BLUE)
for b in ["UI event (onGenerateClick) and the ViewModel function call.",
          "Coroutine scope ownership: the viewModelScope.launch origin annotation on the ViewModel lifeline.",
          "The Repository decision point (\u201ccache fresh?\u201d).",
          "Branch to Room (cache/local path) and branch to the API (remote path).",
          "The mapping-layer invocation (DTO \u2192 Domain \u2192 Entity), annotated on the Repository lifeline in the IO band.",
          "Where caching occurs (write-through upsert to Room).",
          "Flow emission back through the ViewModel StateFlow and UI recomposition.",
          "Thread boundary transitions Main \u2192 IO \u2192 Main, drawn as horizontal bands."]:
    bullet(doc, b)
page_break(doc)

# ---------------------------------------------------------------------------
# 21. ARCHITECTURAL RISK STATEMENT
# ---------------------------------------------------------------------------
H1(doc, "21. Architectural Risk Statement")
P(doc, "Introducing a second data source creates risk. Below we identify risks across the required "
  "categories \u2014 structural, lifecycle, and concurrency \u2014 plus additional professional-grade "
  "risks, and explain how the architecture mitigates each.")

H2(doc, "21.1 Structural Risk")
add_table(doc, ["Risk", "Mitigation"],
    [["Layer bypassing: a ViewModel or Composable calls ApiService or a DAO directly once networking exists, leaking DTO/Room types upward and coupling the UI to the wire format.",
      "The Repository is the only class that holds ApiService and DAO references. The public Repository API is domain-typed only. Code review and the Layer Boundary Rules (\u00a78.5, SD-1\u2013SD-6) forbid any other layer from importing Retrofit, Room, or DTO types."],
     ["Model bleed: DTOs reused as domain or entity types, so a JSON change forces edits across UI, persistence, and tests.",
      "Three separate model families (\u00a717) with mapping isolated in the Repository. A contract change is absorbed inside one mapper."]],
    widths=[3.6, 4.0], font=9)

H2(doc, "21.2 Lifecycle Risk")
add_table(doc, ["Risk", "Mitigation"],
    [["Leaked / orphaned API calls: a long AI request continues after the user navigates away or rotates, wasting work and risking a callback into a dead ViewModel.",
      "All requests run in viewModelScope; onCleared() cancels them via structured concurrency. No GlobalScope. Cancellation propagates to the suspending HTTP call."],
     ["Lost state across configuration change: in-flight loading state or results discarded on rotation.",
      "The ViewModel survives configuration changes and holds StateFlow; the UI re-collects current state on recomposition. Results are cached to Room, so even process death recovers via the Room-backed Flow."]],
    widths=[3.6, 4.0], font=9)

H2(doc, "21.3 Concurrency Risk")
add_table(doc, ["Risk", "Mitigation"],
    [["Main-thread blocking: accidentally performing network/disk work on Main, causing jank or ANR.",
      "withContext(Dispatchers.IO) wraps all network/database/mapping; sensor math uses Dispatchers.Default. Only emissions run on Main (\u00a718)."],
     ["Cache write race: concurrent refreshInsights() calls interleave reads and write-through upserts, producing stale or duplicated cache rows.",
      "Single-source-of-truth write-through with an upsert keyed on a stable LATEST_ID; reads come from the reactive Room Flow. If needed, refreshes are serialized via a Mutex in the Repository so only one network refresh runs at a time."]],
    widths=[3.6, 4.0], font=9)

H2(doc, "21.4 Additional Risks (Beyond Minimum)")
add_table(doc, ["Risk", "Category", "Mitigation"],
    [["Offline / network failure leaves the Insights screen empty or stuck loading.", "Resilience",
      "Repository returns the cached insight on failure (graceful degradation); UI shows a \u201cshowing last result\u201d banner. NFR-4 satisfied."],
     ["Schema drift: backend adds/removes JSON fields, breaking strict parsing.", "Evolvability",
      "DTO fields are nullable with defaults and lenient Gson parsing; unknown fields are ignored; mapping supplies safe defaults."],
     ["Cache staleness: user sees outdated insights after new data is logged.", "Data integrity",
      "Freshness policy (isFresh() TTL) plus an explicit Generate action force a refresh; generatedAt drives invalidation."],
     ["Privacy leak via payload: sending PII to the API.", "Compliance",
      "The Repository constructs the payload from aggregated metrics only \u2014 no name/email/location (\u00a72.4, NFR-5)."],
     ["Migration cost: adding a CachedInsight table breaks existing Room DBs.", "Maintainability",
      "Versioned Room migration adds the new table without touching existing entities; ViewModel/UI unchanged."]],
    widths=[3.0, 1.3, 3.3], font=8.5)

callout(doc, "Evaluation alignment: the Repository is the sole mediator between data sources (\u00a715); "
        "layering is enforced with no bypassing (\u00a78.5); DTO/Domain/Entity are cleanly separated "
        "(\u00a717); thread boundaries are explicit (\u00a718); the ViewModel remains the state owner and "
        "reactive flow integrity is preserved (\u00a715.3\u201315.4); and the sequence diagram is updated "
        "for the dual-source design (\u00a720).", fill="EAF2F8")

print("part B complete")

# ===========================================================================
# SAVE
# ===========================================================================
OUT_NAME = "BTJ-SRS-060326-GP4-brandonGalli.docx"
import os
out_path = os.path.join(os.path.dirname(__file__), OUT_NAME)
doc.save(out_path)
print("SAVED:", out_path)
print("paragraphs:", len(doc.paragraphs), "tables:", len(doc.tables))
