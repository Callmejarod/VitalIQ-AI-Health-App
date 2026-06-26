# -*- coding: utf-8 -*-
"""Build the VitalIQ COMPLIANCE TRACEABILITY MATRIX.

A standalone deliverable that maps every course/assignment requirement to the
exact place in the submitted codebase that satisfies it, and states how. Every
reference is a real file + class/method in the submitted build (paths are
relative to kotlin/app/src/main/java/com/vitaliq/app/ unless noted).

Run:  python build_compliance_traceability.py
Out:  BTJ-ComplianceTraceability-document-06-26-26-brandonGalli.docx
"""
import io, os, sys
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8")

from docx.enum.text import WD_ALIGN_PARAGRAPH as ALIGN

from docx_helpers import (
    new_doc, set_margins, add_table, H1, H2, H3, P, bullet, code_block,
    add_image, page_break, callout, add_toc_field,
    NAVY, BLUE, TEAL, GREY, DARK, WHITE, AMBER,
)

DOCDIR = os.path.dirname(__file__)
OUT = os.path.join(DOCDIR, "BTJ-ComplianceTraceability-document-06-26-26-brandonGalli.docx")

doc = new_doc()
set_margins(doc, 0.9)
C = ALIGN.CENTER

# Column widths reused for the 3-column traceability tables
W = [1.9, 2.5, 2.7]

# ===========================================================================
# COVER
# ===========================================================================
P(doc, "VitalIQ", bold=True, size=34, color=NAVY, align=C, space_after=0)
P(doc, "Compliance Traceability Matrix", bold=True, size=17, color=BLUE, align=C, space_after=14)
P(doc, "Requirement → Implementation → How It Complies", italic=True, size=12, color=GREY, align=C, space_after=18)

add_table(doc,
    ["Field", "Detail"],
    [["Project", "VitalIQ — native Android (Kotlin + Jetpack Compose)"],
     ["Author", "Brandon Galli"],
     ["Course", "DEV322 — Native Android Engineering"],
     ["Document", "Compliance Traceability Matrix"],
     ["Date", "June 26, 2026"],
     ["Build status", "assembleDebug SUCCESSFUL (APK produced; Room KSP code generated)"],
     ["Source root", "kotlin/app/src/main/java/com/vitaliq/app/"]],
    widths=[1.6, 5.2], font=10)

P(doc, "Every reference below points to a real file and class/method in the submitted build. "
       "File paths are relative to kotlin/app/src/main/java/com/vitaliq/app/ unless a different "
       "root is shown. This matrix is the index a reviewer can use to verify each requirement "
       "directly in the codebase.", italic=True, size=10, color=GREY)
page_break(doc)

# ===========================================================================
# TOC
# ===========================================================================
H1(doc, "Table of Contents")
add_toc_field(doc)
page_break(doc)

# ===========================================================================
# 0. RESOLUTION OF THE PRIOR ARCHITECTURAL VIOLATION
# ===========================================================================
H1(doc, "0. Resolution of the Prior Architectural Violation")
P(doc, "The earlier submission was rejected because each ViewModel declared "
       "private val api = RetrofitClient.apiService and called the API directly, making the "
       "ViewModel responsible for data access. That coupling has been removed project-wide.")
add_table(doc,
    ["Concern raised", "Before", "After (submitted build)"],
    [["ViewModel knows the data source",
      "private val api = RetrofitClient.apiService inside each ViewModel.",
      "ViewModels take a Repository INTERFACE in their constructor; no ViewModel imports Retrofit, Room, a DAO, AppDatabase, or okhttp."],
     ["ViewModel performs network calls",
      "load() called api.dashboardSummary() / api.getProfile().",
      "load() calls repo.getDashboardSummary() / repo.getProfile(); the network call lives in DashboardRepositoryImpl."],
     ["“Replace Retrofit with Room → only the Repository changes”",
      "Failed — every ViewModel owned Retrofit calls.",
      "Holds — data-source wiring is centralized in di/ServiceLocator.kt; swapping a source edits only the *RepositoryImpl + one wiring line, no ViewModel."],
     ["SRS vs. code consistency",
      "SRS said Repository is the sole data-access layer; code contradicted it.",
      "Code now matches the SRS; the Final Report and GP6 docs were reconciled to the as-built design."]],
    widths=[1.9, 2.5, 2.7], font=8.5)
callout(doc, "Verification: a repository-wide search of the UI layer "
             "(ui/) for RetrofitClient | AppDatabase | ApiService | okhttp | retrofit2 | "
             "androidx.room | .dao( returns ZERO matches. The only file that references "
             "RetrofitClient.apiService and AppDatabase.getInstance is di/ServiceLocator.kt.")
page_break(doc)

# ===========================================================================
# 1. APPROVED ARCHITECTURAL STACK
# ===========================================================================
H1(doc, "1. Approved Architectural Stack")
add_table(doc,
    ["Requirement", "Where in code", "How it complies"],
    [["Jetpack Compose for UI",
      "ui/screens/**/*Screen.kt; ui/components/*; ui/theme/VitalTheme.kt",
      "All six screens are @Composable functions; MainActivity hosts the Compose tree via setContent { VitalIQTheme { AppNavigation() } }."],
     ["ViewModel for UI state ownership",
      "ui/screens/**/*ViewModel.kt (6 ViewModels)",
      "Each screen has a dedicated androidx.lifecycle.ViewModel that owns the screen's state."],
     ["Coroutines with viewModelScope",
      "Every ViewModel, e.g. DashboardViewModel.load()",
      "All async work is launched in viewModelScope.launch { } and cancelled automatically in onCleared()."],
     ["Repository abstraction layer",
      "data/repository/*Repository.kt (5 interfaces) + *RepositoryImpl.kt",
      "ViewModels depend on the interfaces; the impls are the only classes that touch ApiService + Room DAOs."],
     ["Room for persistence",
      "data/local/AppDatabase.kt; data/local/dao/* (7); data/local/entity/* (7)",
      "@Database(version=2) with 7 DAOs/entities; codegen via KSP (AppDatabase_Impl generated at build)."],
     ["Navigation component with NavHost",
      "navigation/AppNavigation.kt",
      "NavHost with six composable destinations (home, workout, log, insights, profile, history) + bottom navigation."],
     ["Flow / StateFlow for reactive state",
      "Every ViewModel: private MutableStateFlow + public StateFlow",
      "State is exposed as immutable StateFlow<UiState>; screens collect it reactively."],
     ["No extra DI / state / concurrency framework",
      "di/ServiceLocator.kt; app/build.gradle.kts",
      "Manual composition root in plain Kotlin (allowed). No Hilt/Dagger/Koin/@Inject anywhere; no alt state-management or concurrency library."]],
    widths=W, font=8.5)
page_break(doc)

# ===========================================================================
# 2. CORE PLATFORM INTEGRATION
# ===========================================================================
H1(doc, "2. Core Platform Integration")
add_table(doc,
    ["Requirement", "Where in code", "How it complies"],
    [["≥ 2 distinct hardware sensors",
      "ui/screens/workout/WorkoutViewModel.kt — registerSensors()",
      "Registers Sensor.TYPE_ACCELEROMETER and Sensor.TYPE_STEP_COUNTER via SensorManager.registerListener()."],
     ["Sensor data is not decorative / drives observable state",
      "WorkoutViewModel.onSensorChanged(); emitActiveState()",
      "Accelerometer magnitude classifies activity (stationary/walking/mixed/running) and the step counter yields live steps; both flow into WorkoutUiState.Active (StateFlow) and the saved WorkoutDto."],
     ["Sensors run by default (not simulated)",
      "app/build.gradle.kts — SIMULATE_SENSORS",
      "BuildConfig flag defaults to \"false\", so a clean build reads the real accelerometer + step counter; simulation is opt-in for emulators only."],
     ["Sensor runtime permission",
      "AndroidManifest.xml; ui/screens/workout/WorkoutScreen.kt",
      "ACTIVITY_RECOGNITION declared and requested at runtime (rememberLauncherForActivityResult) on API 29+ before a workout starts."],
     ["≥ 1 external network API via coroutines",
      "data/api/ApiService.kt; data/api/RetrofitClient.kt",
      "Retrofit interface of suspend functions (GET/POST/PUT/Multipart) to a FastAPI backend; called from repositories."],
     ["API interaction influences behavior",
      "DashboardRepositoryImpl, InsightsRepositoryImpl, etc.",
      "Dashboard summary, AI insights, history, and logging all come from / post to the API and change the rendered state."],
     ["≥ 2 navigable destinations in a Nav graph",
      "navigation/AppNavigation.kt",
      "Six destinations defined in one NavHost graph."],
     ["Navigation state survives configuration changes",
      "navigation/AppNavigation.kt — rememberNavController()",
      "The back stack is saved/restored across rotation (rememberSaveable internally; saveState/restoreState on tab switches). No manual Intent wiring."],
     ["All network operations off the main thread",
      "data/repository/*RepositoryImpl.kt",
      "Every suspend data method is wrapped in withContext(Dispatchers.IO)."],
     ["Persist domain data via Room through a Repository",
      "*RepositoryImpl.kt → data/local/dao/*",
      "Repositories write API results through to Room (dao.insert) and read from Room on failure; the UI never touches a DAO."]],
    widths=W, font=8.5)
page_break(doc)

# ===========================================================================
# 3. ARCHITECTURE REQUIREMENTS
# ===========================================================================
H1(doc, "3. Architecture Requirements")
add_table(doc,
    ["Requirement", "Where in code", "How it complies"],
    [["Separate UI / ViewModel / Repository",
      "ui/screens/** ; data/repository/** ; data/local + data/api",
      "Three distinct packages; the UI talks only to the ViewModel, the ViewModel only to a Repository interface, the Repository to the data sources."],
     ["ViewModel is the single source of truth",
      "Every ViewModel: _uiState (MutableStateFlow) / uiState (StateFlow)",
      "Screen state is owned by the ViewModel and exposed read-only; the UI renders purely from it."],
     ["Activity/Composable owns no long-lived state",
      "MainActivity.kt; *Screen.kt",
      "MainActivity only hosts Compose; screens keep only ephemeral UI state (text-field drafts, isRefreshing) — all domain state lives in the ViewModel."],
     ["All data flows through a Repository before the ViewModel",
      "*ViewModel constructors take *Repository interfaces",
      "ViewModels receive data exclusively from repository methods (e.g. repo.listWorkouts())."],
     ["No direct DB/network access from UI or ViewModel",
      "Verified across ui/ (grep) ; di/ServiceLocator.kt",
      "Zero Retrofit/Room/DAO references in any screen or ViewModel; the only data-source construction is in ServiceLocator."],
     ["UI updates reactively to state changes",
      "*Screen.kt — collectAsStateWithLifecycle()",
      "Each screen collects the StateFlow with collectAsStateWithLifecycle() and recomposes on emission."],
     ["Handle configuration changes without data loss",
      "ViewModel + StateFlow; rememberNavController()",
      "The ViewModel survives rotation (state retained); navigation back stack is preserved. Verified by lifecycle logs (Activity recreated, ViewModel NOT cleared)."],
     ["Lifecycle logging, consistent tag, Logcat-visible",
      "MainActivity.kt; every ViewModel; VitalIQApp.kt",
      "Single tag \"VitalIQ\": MainActivity logs onCreate/onResume/onPause/onStop/onDestroy; each ViewModel logs init + onCleared; VitalIQApp logs onCreate. Filter Logcat by \"VitalIQ\" to demonstrate."]],
    widths=W, font=8.5)
P(doc, "Representative evidence — ViewModel depends only on the Repository interface and logs its lifecycle:",
  bold=True, size=9.5, color=BLUE)
code_block(doc,
    "class DashboardViewModel(private val repo: DashboardRepository) : ViewModel() {\n"
    "    private val _uiState = MutableStateFlow<DashboardUiState>(DashboardUiState.Loading)\n"
    "    val uiState: StateFlow<DashboardUiState> = _uiState\n"
    "    init { Log.d(\"VitalIQ\", \"DashboardViewModel created\") }\n"
    "    fun load() = viewModelScope.launch {\n"
    "        try { /* repo.getDashboardSummary(); repo.getProfile() */ }\n"
    "        catch (e: Exception) { _uiState.value = Error(...) }\n"
    "    }\n"
    "    override fun onCleared() { Log.d(\"VitalIQ\", \"DashboardViewModel cleared\") }\n"
    "}")
page_break(doc)

# ===========================================================================
# 4. CONCURRENCY AND STATE DISCIPLINE
# ===========================================================================
H1(doc, "4. Concurrency and State Discipline")
add_table(doc,
    ["Requirement", "Where in code", "How it complies"],
    [["All disk and network IO off the main thread",
      "*RepositoryImpl.kt; ProfileViewModel.importFile()",
      "Repository IO uses withContext(Dispatchers.IO); the profile file read also runs in withContext(Dispatchers.IO)."],
     ["No blocking calls on the main thread",
      "WorkoutViewModel timer; all ViewModels",
      "The workout timer uses kotlinx.coroutines.delay (suspending), not Thread.sleep; no blocking I/O on Main."],
     ["No runBlocking on the main thread",
      "Whole module (static scan)",
      "runBlocking does not appear anywhere in the codebase."],
     ["Appropriate dispatchers",
      "*RepositoryImpl.kt",
      "Network + Room run on Dispatchers.IO; UI state updates run on the Main-confined viewModelScope."],
     ["IO on Dispatchers.IO unless justified otherwise",
      "*RepositoryImpl.kt (withContext(Dispatchers.IO))",
      "Every network/disk operation is on Dispatchers.IO; no undocumented alternative."],
     ["Long-running operations don't block rendering",
      "viewModelScope.launch in every ViewModel",
      "Loads/saves are suspend functions on background dispatchers; the UI shows Loading/Submitting states meanwhile."],
     ["Thread boundaries identifiable in the architecture diagram",
      "documents/diagrams/arch_layered.png; build_diagrams.py",
      "The layered diagram marks the Main→IO boundary (withContext(Dispatchers.IO)) and Main/IO/Default bands; activity & sequence diagrams show the same boundaries."]],
    widths=W, font=8.5)
page_break(doc)

# ===========================================================================
# 5. RESILIENCE AND PROFESSIONAL BEHAVIOR
# ===========================================================================
H1(doc, "5. Resilience and Professional Behavior")
add_table(doc,
    ["Requirement", "Where in code", "How it complies"],
    [["Gracefully handle network failure with a meaningful UI state",
      "*UiState.Error; DashboardScreen/HistoryScreen/ProfileScreen/InsightsScreen",
      "Each loading flow has a sealed Error state rendered as a visible message; failures are caught in the ViewModel and mapped to Error."],
     ["Visible UI state + recovery path",
      "DashboardScreen/HistoryScreen/ProfileScreen/InsightsScreen Error branch",
      "Error UI shows a warning + message and a Retry button wired to viewModel.load(); Insights also shows a degraded local-estimate banner (usedFallback)."],
     ["Avoid crashes (sensor / rotation / network)",
      "WorkoutViewModel.onCleared(); try/catch in repos; rememberNavController",
      "Sensor listeners + timer cancelled in onCleared(); repository calls are guarded; rotation retains the ViewModel; missing step sensor handled (stepCounterAvailable)."],
     ["Network failure demonstrable (airplane mode / API failure)",
      "*RepositoryImpl.kt offline fallback",
      "With the backend unreachable, repositories catch the exception and return cached Room data; if no cache, the ViewModel surfaces an Error state. Both are demonstrable by toggling airplane mode."],
     ["Error states are visible and intentional",
      "Sealed UiState hierarchies (Loading/Success/Error/Empty/…)",
      "Errors are first-class states, not silent failures; the UI renders a distinct, intentional error layout."]],
    widths=W, font=8.5)
P(doc, "Representative evidence — repository network-first with Room offline fallback:",
  bold=True, size=9.5, color=BLUE)
code_block(doc,
    "override suspend fun listWorkouts(): List<WorkoutDto> = withContext(Dispatchers.IO) {\n"
    "    try {\n"
    "        val workouts = api.listWorkouts()\n"
    "        dao.insertAll(workouts.map { it.toEntity() })   // write-through to Room\n"
    "        workouts\n"
    "    } catch (e: Exception) {\n"
    "        val cached = dao.getAll()                        // offline fallback\n"
    "        if (cached.isEmpty()) throw e else cached.map { it.toDto() }\n"
    "    }\n"
    "}")
page_break(doc)

# ===========================================================================
# 6. ARCHITECTURAL DOCUMENTATION & DEFENSE
# ===========================================================================
H1(doc, "6. Architectural Documentation & Defense")
add_table(doc,
    ["Requirement", "Where", "How it complies"],
    [["Documentation matches the submitted build",
      "documents/build_final_report.py; build_gp6.py; build_srs.py; kotlin/README.md",
      "All architecture docs were reconciled to the as-built Repository + Room design and regenerated; the legacy memory/PRD.md is marked superseded."],
     ["Architectural claims verifiable in the codebase",
      "This matrix (every row cites a file + symbol)",
      "Each documented claim maps to a concrete file/class/method that a reviewer can open."],
     ["Diagrams include class names",
      "documents/diagrams/arch_layered.png; seq_insights_week5.png",
      "The layered diagram labels real classes (DashboardViewModel, *Repository, ApiService, AppDatabase, ServiceLocator); the sequence diagram names the same participants."],
     ["Layered architecture diagram (UI/VM/Repo/sources/thread boundaries)",
      "documents/diagrams/arch_layered.png (Figure 3.1 in the Final Report)",
      "Shows all four layers, the composition root, both data sources, the Main→IO boundary, and Main/IO/Default dispatcher bands — reflecting the implemented design, not an idealized one."],
     ["Diagram indicates Main / IO / Default dispatchers",
      "documents/diagrams/arch_layered.png; act_*.png",
      "Explicit MAIN THREAD and IO DISPATCHER bands; the boundary is annotated withContext(Dispatchers.IO); sensor classification noted as Default-eligible."]],
    widths=W, font=8.5)

H2(doc, "6.1 Where the “Why” Answers Live in the Code")
P(doc, "The required defense questions map to concrete implementation decisions:")
add_table(doc,
    ["Question", "Concrete answer in this codebase"],
    [["Why a ViewModel exists",
      "State must survive rotation: the ViewModel + StateFlow retain DashboardUiState/WorkoutUiState across config changes (lifecycle logs show the Activity recreates while the ViewModel does not)."],
     ["Why coroutines instead of raw threads",
      "viewModelScope ties work to the ViewModel lifecycle (auto-cancel in onCleared); withContext(Dispatchers.IO) moves IO off Main without manual thread/handler management."],
     ["Why Room instead of manual SQLite",
      "AppDatabase + DAOs give compile-time-checked queries and generated implementations (KSP) instead of hand-written SQL/cursor code; entities map cleanly to DTOs via toEntity()/toDto()."],
     ["Why reactive state collection",
      "StateFlow + collectAsStateWithLifecycle() make the UI a pure function of state and stop collection when not in the STARTED lifecycle, avoiding wasted work/leaks."],
     ["Why a navigation graph instead of manual intents",
      "A single NavHost owns the back stack and survives rotation (rememberNavController), versus brittle startActivity/Intent wiring and manual state restoration."]],
    widths=[1.9, 5.2], font=8.5)
page_break(doc)

# ===========================================================================
# 7. DEMONSTRATION CRITERIA
# ===========================================================================
H1(doc, "7. Demonstration Criteria — Where Each Is Supported")
P(doc, "Each required demonstration item is backed by the code below; the live actions are "
       "performed during the presentation (the FastAPI backend must be running for API-backed flows).")
add_table(doc,
    ["Demonstration item", "Backed by", "How to show it"],
    [["Sensor interaction",
      "WorkoutViewModel (accelerometer + step counter)",
      "Start a workout; live steps + activity classification update from real sensors."],
     ["API interaction",
      "ApiService via repositories",
      "Open Dashboard/Insights; data loads from the backend through the repository."],
     ["Local persistence",
      "Room AppDatabase via repositories",
      "Load data online (caches to Room), then read it back; inspect with App Inspection → Database."],
     ["Reactive UI updates",
      "StateFlow + collectAsStateWithLifecycle",
      "Log an entry / generate insights; the screen recomposes from new state immediately."],
     ["Rotation survival",
      "ViewModel + StateFlow + rememberNavController",
      "Rotate during an active workout; state + nav position are retained (Logcat \"VitalIQ\" shows Activity recreate, ViewModel retained)."],
     ["Graceful simulated network failure",
      "Repository offline fallback + Error UI states",
      "Enable airplane mode; reads fall back to cached Room data, or an Error state with Retry appears — no crash."]],
    widths=[1.7, 2.3, 3.1], font=8.5)
page_break(doc)

# ===========================================================================
# 8. THREAD / DISPATCHER BOUNDARY SUMMARY
# ===========================================================================
H1(doc, "8. Thread & Dispatcher Boundary Summary")
add_image(doc, "arch_layered.png", width=6.7,
          caption="Layered architecture with Main→IO thread boundary and dispatcher bands (as-built).")
add_table(doc,
    ["Operation", "Thread / Dispatcher", "Where"],
    [["UI rendering / recomposition", "Main", "*Screen.kt (Compose)"],
     ["ViewModel state updates (StateFlow)", "Main (viewModelScope)", "*ViewModel.kt"],
     ["Composition-root wiring (lazy build)", "Main (caller)", "di/ServiceLocator.kt"],
     ["Network calls (Retrofit suspend)", "IO", "*RepositoryImpl.kt withContext(Dispatchers.IO)"],
     ["Room reads/writes (DAO)", "IO", "*RepositoryImpl.kt → data/local/dao/*"],
     ["Profile file read (ContentResolver)", "IO", "ProfileViewModel.importFile()"],
     ["Sensor activity classification", "Main callback (Default-eligible, CPU-light)", "WorkoutViewModel.onSensorChanged()"]],
    widths=[2.4, 2.3, 2.4], font=8.5)
page_break(doc)

# ===========================================================================
# 9. SUMMARY SCORECARD
# ===========================================================================
H1(doc, "9. Summary Scorecard")
add_table(doc,
    ["Requirement area", "Status"],
    [["Approved architectural stack (Compose/VM/coroutines/Repository/Room/Nav/StateFlow)", "Met"],
     ["No prohibited DI / state / concurrency framework", "Met (manual ServiceLocator)"],
     ["Two hardware sensors driving observable state", "Met"],
     ["External network API via coroutines, off main thread", "Met"],
     ["≥ 2 nav destinations; survive configuration change", "Met"],
     ["Persist via Room through a Repository", "Met"],
     ["UI / ViewModel / Repository separation; no direct data access in UI/VM", "Met"],
     ["ViewModel single source of truth; reactive UI", "Met"],
     ["Configuration-change survival; lifecycle logging (tag \"VitalIQ\")", "Met"],
     ["Concurrency discipline (IO off Main, no runBlocking, Dispatchers.IO)", "Met"],
     ["Graceful network-failure handling with visible state + recovery", "Met"],
     ["Documentation matches build; layered diagram with class names + thread boundaries", "Met"],
     ["Build compiles & runs (assembleDebug SUCCESSFUL; Room KSP generated)", "Met"]],
    widths=[5.3, 1.9], font=9)
callout(doc, "The instructor's specific architectural violation — ViewModels instantiating and "
             "invoking RetrofitClient.apiService — is fully resolved: data access is mediated by the "
             "Repository layer, ViewModels are data-source agnostic, and the documentation now "
             "matches the implemented architecture.")

# ===========================================================================
# SAVE
# ===========================================================================
doc.save(OUT)
print("SAVED:", OUT)
print("paragraphs:", len(doc.paragraphs), "tables:", len(doc.tables),
      "images:", len(doc.inline_shapes))
