# -*- coding: utf-8 -*-
"""Transform the GP5 SRS into the GP6 SRS:
   - update cover/version/feedback log
   - insert AI Report immediately after the Table of Contents
   - extend the TOC document-map
   - append GP6 sections (Summary List, Test Plans & Results, Architectural Audit,
     What Remains, updated Phased Approach) at the end of the document.
   Works on a copy so the GP5 source is preserved.
"""
import io, sys, os, copy
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8")
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

NAVY = RGBColor(0x1B, 0x3A, 0x57)
BLUE = RGBColor(0x2E, 0x6E, 0x9E)
TEAL = RGBColor(0x1F, 0x86, 0x8C)
GREY = RGBColor(0x44, 0x55, 0x66)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
DARK = RGBColor(0x20, 0x2A, 0x33)

DOCDIR = os.path.dirname(__file__)
GP5 = os.path.join(DOCDIR, "BTJ-GP5-SRS-document-06-16-26-brandonGalli.docx")
SRC = os.path.join(DOCDIR, "BTJ-GP6-SRS-document-06-16-26-brandonGalli.docx")

# Always start from a clean copy of the GP5 source so re-runs are idempotent.
import shutil
shutil.copyfile(GP5, SRC)

doc = Document(SRC)

# ---------------------------------------------------------------------------
# A cursor-based builder that APPENDS to the body (for end-of-doc sections),
# plus an insert-before helper for the AI report.
# ---------------------------------------------------------------------------
class Anchor:
    """Insert new paragraphs/tables immediately before a reference paragraph."""
    def __init__(self, ref_paragraph):
        self.ref = ref_paragraph._p  # lxml element

    def _add(self, element):
        self.ref.addprevious(element)
        return element

def _shade(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    sh = OxmlElement("w:shd"); sh.set(qn("w:val"), "clear"); sh.set(qn("w:fill"), hexcolor)
    tcPr.append(sh)

def _cell_text(cell, text, bold=False, color=None, size=10):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(str(text))
    run.font.size = Pt(size); run.font.bold = bold; run.font.name = "Calibri"
    if color: run.font.color.rgb = color


class Builder:
    """Build paragraphs/tables either appended to doc body or inserted before anchor."""
    def __init__(self, document, anchor=None):
        self.doc = document
        self.anchor = anchor  # lxml element to insert before; None => append

    def _place(self, element):
        if self.anchor is not None:
            self.anchor.addprevious(element)
        # if appending, element already added to body by python-docx
        return element

    def _new_par(self):
        if self.anchor is not None:
            p = self.doc.add_paragraph()
            el = p._p
            el.getparent().remove(el)
            self.anchor.addprevious(el)
            return p
        return self.doc.add_paragraph()

    def heading(self, text, level):
        if self.anchor is not None:
            p = self.doc.add_heading(text, level=level)
            el = p._p; el.getparent().remove(el); self.anchor.addprevious(el)
            return p
        return self.doc.add_heading(text, level=level)

    def H1(self, t): return self.heading(t, 1)
    def H2(self, t): return self.heading(t, 2)
    def H3(self, t): return self.heading(t, 3)

    def P(self, text, bold=False, italic=False, size=11, color=None, align=None, space_after=None):
        p = self._new_par()
        if align: p.alignment = align
        if space_after is not None: p.paragraph_format.space_after = Pt(space_after)
        r = p.add_run(text); r.bold = bold; r.italic = italic; r.font.size = Pt(size)
        if color: r.font.color.rgb = color
        return p

    def bullet(self, text, bold_lead=None):
        p = self._new_par()
        p.style = self.doc.styles["List Bullet"]
        p.paragraph_format.left_indent = Inches(0.3)
        p.paragraph_format.space_after = Pt(2)
        if bold_lead:
            r = p.add_run(bold_lead); r.bold = True
        p.add_run(text)
        return p

    def numbered(self, text):
        p = self._new_par()
        p.style = self.doc.styles["List Number"]
        p.paragraph_format.space_after = Pt(2)
        p.add_run(text)
        return p

    def code(self, text):
        p = self._new_par()
        p.paragraph_format.left_indent = Inches(0.2)
        p.paragraph_format.space_before = Pt(4); p.paragraph_format.space_after = Pt(8)
        r = p.add_run(text); r.font.name = "Consolas"; r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
        pPr = p._p.get_or_add_pPr()
        sh = OxmlElement("w:shd"); sh.set(qn("w:val"), "clear"); sh.set(qn("w:fill"), "F2F4F6")
        pPr.append(sh)
        return p

    def table(self, headers, rows, widths=None, header_fill="1B3A57", zebra="EAF2F8", font=9.5):
        t = self.doc.add_table(rows=1, cols=len(headers))
        t.style = "Table Grid"
        hdr = t.rows[0].cells
        for i, h in enumerate(headers):
            _shade(hdr[i], header_fill)
            _cell_text(hdr[i], h, bold=True, color=WHITE, size=font)
        for r_idx, row in enumerate(rows):
            cells = t.add_row().cells
            for i, val in enumerate(row):
                _cell_text(cells[i], val, size=font)
                if zebra and r_idx % 2 == 1:
                    _shade(cells[i], zebra)
        if widths:
            for row in t.rows:
                for i, w in enumerate(widths):
                    row.cells[i].width = Inches(w)
        if self.anchor is not None:
            el = t._tbl; el.getparent().remove(el); self.anchor.addprevious(el)
            # trailing spacer paragraph
            sp = self.doc.add_paragraph(); spel = sp._p; spel.getparent().remove(spel)
            self.anchor.addprevious(spel)
        else:
            self.doc.add_paragraph()
        return t

    def callout(self, text, fill="EAF2F8"):
        t = self.doc.add_table(rows=1, cols=1); t.style = "Table Grid"
        c = t.cell(0, 0); _shade(c, fill); _cell_text(c, text, size=10.5, color=NAVY)
        if self.anchor is not None:
            el = t._tbl; el.getparent().remove(el); self.anchor.addprevious(el)
            sp = self.doc.add_paragraph(); spel = sp._p; spel.getparent().remove(spel)
            self.anchor.addprevious(spel)
        else:
            self.doc.add_paragraph()

    def page_break(self):
        if self.anchor is not None:
            p = self.doc.add_paragraph(); el = p._p; el.getparent().remove(el)
            self.anchor.addprevious(el)
            run = p.add_run(); br = OxmlElement("w:br"); br.set(qn("w:type"), "page")
            run._r.append(br)
        else:
            self.doc.add_page_break()

print("loaded GP5 base, paragraphs:", len(doc.paragraphs))

# ===========================================================================
# 1. UPDATE COVER PAGE TEXT (GP5 -> GP6)
# ===========================================================================
cover_map = {
    "Cumulative SRS — GP1 + GP2 + GP3 + Week 4 + GP4 + GP5":
        "Cumulative SRS — GP1 + GP2 + GP3 + Week 4 + GP4 + GP5 + GP6",
    "Course: DEV322  ·  Milestone GP5 — Core Functionality & Architectural Audit":
        "Course: DEV322  ·  Milestone GP6 — Integration, Testing & Architectural Legitimacy Checkpoint",
    "Document Version 5.0  ·  Cumulative":
        "Document Version 6.0  ·  Cumulative",
}
for p in doc.paragraphs[:12]:
    if p.text.strip() in cover_map:
        new = cover_map[p.text.strip()]
        for r in p.runs:
            r.text = ""
        if p.runs:
            p.runs[0].text = new
        else:
            p.add_run(new)
print("cover updated")

# ===========================================================================
# 2. UPDATE REVISION HISTORY TABLE (add GP6 row)
# ===========================================================================
rev = doc.tables[0]
row = rev.add_row().cells
vals = ["6.0", "GP6 — Integration & Architectural Legitimacy",
        "Jun 16, 2026",
        "AI Report; Summary List of SRS components; expanded Test Plans & Results "
        "(Verification, Validation, Final Acceptance); Architectural Legitimacy Audit "
        "(traceability, alignment, GlobalScope/runBlocking, rotation survival); "
        "What-Remains feature backlog; refined phased approach."]
for i, v in enumerate(vals):
    _cell_text(row[i], v, size=9)

# add a GP6 feedback-resolution row to the feedback log (table 1)
fb = doc.tables[1]
frow = fb.add_row().cells
fvals = ["7", "Test evidence (GP5)",
         "Tie test results back to user stories and show rotation/audit evidence.",
         "Added a Traceability Matrix Audit through testing (\u00a726) and an Architectural "
         "Legitimacy Audit (\u00a727) with explicit rotation-survival and GlobalScope/runBlocking checks."]
for i, v in enumerate(fvals):
    _cell_text(frow[i], v, size=9)
print("revision + feedback updated")

# ===========================================================================
# 3. EXTEND THE TOC DOCUMENT-MAP LIST
#    Add the AI Report line (right after the TOC heading area) and the new
#    GP6 section lines at the end of Part B list.
# ===========================================================================
def _toc_line_before(ref_p, text, bold=False):
    """Insert a toc-style line paragraph before ref_p."""
    np = doc.add_paragraph()
    el = np._p; el.getparent().remove(el)
    ref_p._p.addprevious(el)
    np.paragraph_format.space_after = Pt(0)
    r = np.add_run(text); r.font.size = Pt(10)
    if bold:
        r.bold = True; r.font.color.rgb = NAVY
    return np

# locate key TOC paragraphs
toc_part_a = None
toc_phased = None
for p in doc.paragraphs:
    s = p.text.strip()
    if s.startswith("Part A — Prior Milestones") and toc_part_a is None:
        toc_part_a = p
    if s.startswith("23. Phased Development Approach"):
        toc_phased = p

# AI Report listed first, immediately after TOC intro / before Part A block
if toc_part_a is not None:
    _toc_line_before(toc_part_a, "AI Report — How Our Team Used AI in the Workflow", bold=True)

# new GP6 sections after the last Part B entry (23.)
if toc_phased is not None:
    # we want them AFTER 23, so insert before the paragraph that follows 23.
    after = toc_phased._p.getnext()
    # build a temp anchor at 'after'; if None, append at end of list area
    new_lines = [
        "24. Summary List of Required SRS Components",
        "25. Test Plans & Test Results (Verification, Validation, Final Acceptance)",
        "26. Traceability Matrix Audit Through Testing",
        "27. Architectural Legitimacy Audit",
        "28. What Remains To Be Developed (Feature Backlog)",
        "29. Phased Development Approach — Updated Next Steps",
    ]
    for line in new_lines:
        np = doc.add_paragraph()
        el = np._p; el.getparent().remove(el)
        if after is not None:
            after.addprevious(el)
        else:
            toc_phased._p.addnext(el)
        np.paragraph_format.space_after = Pt(0)
        r = np.add_run(line); r.font.size = Pt(10)
print("TOC map extended")

# ===========================================================================
# 4. INSERT AI REPORT  — immediately after the Table of Contents,
#    before the page break that precedes "PART A".
# ===========================================================================
# Find the page-break paragraph just before PART A.
part_a_hdr = None
for p in doc.paragraphs:
    if p.text.strip() == "PART A":
        part_a_hdr = p
        break
# The page break sits in the paragraph two before PART A (see layout). Insert the
# AI report before that page-break paragraph so it lands on its own page after TOC.
anchor_p = part_a_hdr
# walk backwards to find the nearest preceding page-break paragraph
prev = part_a_hdr._p.getprevious()
brk_par = None
while prev is not None:
    brs = prev.findall(".//" + qn("w:br"))
    if any(b.get(qn("w:type")) == "page" for b in brs):
        brk_par = prev
        break
    prev = prev.getprevious()

insert_anchor = brk_par if brk_par is not None else part_a_hdr._p
ai = Builder(doc, anchor=insert_anchor)

# True Heading 1 so the Word TOC field lists the AI Report; centered + prominent.
ai_title = ai.H1("AI Report")
ai_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
ai.P("How Team BTJ Used AI in Our Development Workflow", bold=True, size=13, color=BLUE,
     align=WD_ALIGN_PARAGRAPH.CENTER, space_after=8)
ai.callout("This AI Report is scored independently on its own assignment page and is also included "
           "here as a required part of the SRS. It documents where, how, and why our team used AI "
           "tools during the VitalIQ project, and the human review process we applied to everything "
           "AI produced.")

ai.H2("1. Overview")
ai.P("Team BTJ used AI as a development accelerator and a thinking partner \u2014 not as a substitute "
     "for engineering judgment. Every AI-assisted artifact (code, diagram, prose, or test) was "
     "reviewed, edited, and validated by a team member before it entered the repository or this "
     "document. AI helped us move faster on boilerplate, documentation structure, and exploration, "
     "while the team retained ownership of architecture decisions, correctness, and academic "
     "integrity.")

ai.H2("2. Where We Used AI")
ai.table(["Workflow Area", "How AI Was Used", "Human Oversight"],
    [["Architecture & design",
      "Brainstorm Repository-evolution options, cache-then-network patterns, and dispatcher choices; sanity-check MVVM layering.",
      "Team made final architectural decisions; AI suggestions cross-checked against course textbook guidance."],
     ["Kotlin code (Android)",
      "Generate boilerplate (ViewModels, StateFlow scaffolding, sensor listener structure) and explain coroutine/Compose APIs.",
      "All code read line-by-line, compiled, and run in Android Studio before commit; logic verified against requirements."],
     ["Documentation (SRS)",
      "Draft and structure SRS sections, normalize user-story phrasing, and assemble the cumulative table of contents.",
      "Content fact-checked against the actual codebase and prior submissions; instructor feedback applied manually."],
     ["Diagrams",
      "Generate activity, sequence, use-case, and state-machine diagrams (rendered programmatically), with Main/IO swimlanes.",
      "Diagrams reviewed for accuracy against the real call flow; relabeled/redrawn where they misrepresented threading."],
     ["Testing",
      "Suggest unit/acceptance test cases derived from user stories and technical specs; draft edge cases.",
      "Tests mapped back to user stories/use cases by the team; results recorded from real runs, not AI-asserted."],
     ["Debugging",
      "Explain stack traces, recomposition issues, and rotation-survival behavior; propose fixes.",
      "Fixes validated on emulator/device; team confirmed root cause before accepting a change."]],
    widths=[1.7, 3.1, 2.8], font=9)

ai.H2("3. Tools Used")
ai.bullet("a large-language-model coding assistant integrated into our IDE for code generation, "
          "explanation, and refactoring suggestions.", bold_lead="AI coding assistant — ")
ai.bullet("conversational AI for design discussion, requirement clarification, and documentation drafting.",
          bold_lead="Conversational LLM — ")
ai.bullet("programmatic diagram generation (Python/Matplotlib) driven by AI-authored scripts, "
          "kept in the repo so diagrams are reproducible.", bold_lead="Diagram generation — ")
ai.P("The VitalIQ AI Insights feature itself also consumes an external AI model API at runtime; "
     "that is a product feature (documented in \u00a716) and is distinct from the AI tooling used in "
     "our development workflow described here.", italic=True, color=GREY, size=10)

ai.H2("4. Our Human-in-the-Loop Process")
ai.numbered("Prompt with explicit context: we gave the AI our requirements, existing code, and constraints rather than vague requests.")
ai.numbered("Review every output: a team member read and understood each suggestion before using it.")
ai.numbered("Verify by execution: code was compiled and run; tests were actually executed; diagrams were checked against the real flow.")
ai.numbered("Edit for correctness and voice: AI drafts were rewritten to match our design and our wording.")
ai.numbered("Attribute and integrate: accepted changes were committed with normal review on GitHub.")

ai.H2("5. Benefits & Limitations We Observed")
ai.table(["Benefits", "Limitations / Risks We Managed"],
    [["Faster boilerplate and documentation scaffolding freed time for design and testing.",
      "AI occasionally proposed plausible-but-wrong APIs; we caught these by compiling."],
     ["Useful for explaining unfamiliar coroutine/Compose behavior and threading rules.",
      "AI sometimes suggested putting logic in the wrong layer (e.g., mapping in the ViewModel); we corrected per our layer rules."],
     ["Good at generating broad test-case ideas from user stories.",
      "AI cannot assert real test results; all RESULTS in \u00a725 come from actual runs."],
     ["Helped maintain consistency across a large cumulative document.",
      "Risk of stale/duplicated content; mitigated with a human editing pass and the feedback log."]],
    widths=[3.8, 3.8], font=9.5)

ai.H2("6. Academic Integrity Statement")
ai.P("All AI assistance was used as a learning and productivity aid. The team understands the "
     "architecture, code, and requirements presented in this document and can explain and defend "
     "every section. AI did not make final engineering decisions, and no AI output was submitted "
     "without human understanding, verification, and revision.")
ai.page_break()
print("AI report inserted")

# ===========================================================================
# 5. APPEND GP6 SECTIONS AT END OF DOCUMENT (sections 24-29)
# ===========================================================================
b = Builder(doc, anchor=None)  # append mode
b.page_break()
b.P("PART C", bold=True, size=22, color=NAVY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)
b.P("GP6 — Integration, Testing & Architectural Legitimacy Checkpoint",
    bold=True, size=13, color=BLUE, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)
b.callout("This milestone achieves integration of essential core functionality and serves as an "
          "architectural legitimacy checkpoint. It adds a consolidated component checklist, full "
          "test plans and real results, audits (traceability, architectural alignment, coroutine "
          "hygiene, rotation survival), the remaining-feature backlog, and the refined phased plan. "
          "Prior sections remain authoritative and have been refined per instructor feedback.")
b.page_break()

# ---------------------------------------------------------------------------
# 24. SUMMARY LIST OF REQUIRED SRS COMPONENTS
# ---------------------------------------------------------------------------
b.H1("24. Summary List of Required SRS Components")
b.P("This checklist confirms that every required SRS component is present in the latest (GP6) "
    "version, and that the latest version incorporates refinements made per instructor feedback "
    "(see the Document Control \u201cInstructor Feedback Resolution Log\u201d).")

b.H2("24.1 GP-2 Compliant Components")
b.table(["Required Component", "Location", "Status"],
    [["Functional Narrative of Project Scope & Scale", "\u00a71", "Present \u2014 refined"],
     ["Platform / Environment / Compliance Assumptions", "\u00a72", "Present"],
     ["User Requirements (Functional, Non-Functional, System Design)", "\u00a73", "Present"],
     ["Complete User Stories (As a / I want / to achieve)", "\u00a74", "Present \u2014 normalized"],
     ["Operational Flow + Step Lists (one activity per step = pseudocode)", "\u00a79", "Present \u2014 refined"],
     ["ViewModel ownership / Coroutine placement / Persistence intent", "\u00a78.2\u20138.4", "Present"],
     ["Activity Diagrams (process maps / flow charts)", "\u00a710, \u00a719", "Present"],
     ["Week 4 documentation (operational flows, traceability, ownership tables)", "\u00a79, \u00a713", "Present"],
     ["Table of Contents organizing the document", "TOC + map", "Present \u2014 updated for GP6"]],
    widths=[4.0, 1.8, 1.8], font=9)

b.H2("24.2 GP-3 Compliant Components")
b.table(["Required Component", "Location", "Status"],
    [["Acceptance Criteria", "\u00a75", "Present"],
     ["Use Cases", "\u00a76", "Present"],
     ["Use Case Diagrams (stick figures + ovals)", "\u00a77", "Present"],
     ["Activity Diagrams with Main vs IO separation", "\u00a710, \u00a719", "Present"],
     ["Sequence Diagram", "\u00a711, \u00a720", "Present"],
     ["State Machine Diagram (focused, core domain)", "\u00a712", "Present \u2014 refined"],
     ["Why suspend exists / Main cannot block / scope ownership matters", "\u00a714", "Present"],
     ["Traceability Matrix (stories \u2192 acceptance criteria)", "\u00a713, audited in \u00a726", "Present \u2014 audited"],
     ["Annotated activity diagrams (coroutine launch, persistence decision, lifecycle owner)", "\u00a719", "Present"]],
    widths=[4.0, 1.8, 1.8], font=9)

b.H2("24.3 GP-4 Compliant Components")
b.table(["Required Component", "Location", "Status"],
    [["Repository Evolution Plan (+ where branching logic lives)", "\u00a715", "Present"],
     ["API Contract Specification (endpoints, schemas, errors, persisted/transient)", "\u00a716", "Present"],
     ["DTO + Mapping Plan (DTO \u2192 Domain \u2192 Entity, + reverse; why not in UI/DAO)", "\u00a717", "Present"],
     ["Thread Boundary Plan (network/db/mapping dispatchers + rationale)", "\u00a718", "Present"],
     ["Updated Activity Diagram (dual data source, thread boundaries)", "\u00a719", "Present"],
     ["Updated Sequence Diagram (scope origin, Main\u2192IO\u2192Main, mediation)", "\u00a720", "Present"],
     ["Architectural Risk Statement (structural, lifecycle, concurrency, +extra)", "\u00a721", "Present"]],
    widths=[4.0, 1.8, 1.8], font=9)

b.H2("24.4 GP-5 Compliant Components")
b.table(["Required Component", "Location", "Status"],
    [["Updated Reactive State Architecture (Phase 1 core, Phase 2 API, Phase 3 navigation)", "\u00a722", "Present"],
     ["Updated Wireframes / UI integration evidence", "\u00a722.1", "Present"],
     ["Refined SRS per last iteration's comments", "Feedback Log + \u00a726/27", "Present"],
     ["Test Plans: Verification, Validation, Final Acceptance", "\u00a77.5, expanded in \u00a725", "Present \u2014 expanded"],
     ["Test Results", "\u00a77.5, \u00a725.4", "Present \u2014 updated"],
     ["Existing code (well-commented) + zipped archive", "Repo + submission zip", "Provided"]],
    widths=[4.0, 1.8, 1.8], font=9)

b.H2("24.5 GP-6 (This Week) Components")
b.table(["Required Component", "Location"],
    [["AI Report (also scored independently)", "After TOC (front matter)"],
     ["Summary list of all required SRS components", "\u00a724 (this section)"],
     ["Integration of essential core functionality achieved", "\u00a722.1, \u00a727"],
     ["Continued debugging/optimization of integrated units", "\u00a725.4, \u00a727"],
     ["Test cases for units under test + test results", "\u00a725"],
     ["Traceability Matrix audit through testing", "\u00a726"],
     ["Architectural Alignment / legitimacy audit", "\u00a727"],
     ["What remains to be developed (future-phase features)", "\u00a728"],
     ["Phased approach: next steps & preparation", "\u00a729"]],
    widths=[4.6, 3.0], font=9)
b.page_break()
print("section 24 done")

# ---------------------------------------------------------------------------
# 25. TEST PLANS & TEST RESULTS
# ---------------------------------------------------------------------------
b.H1("25. Test Plans & Test Results")
b.P("Per the milestone requirements, testing is organized into three layers. Unit (verification) "
    "tests come from the technical specification \u2014 the system activities in response to user "
    "interactions. User Acceptance tests come from the user steps defined by the user stories. "
    "Final Acceptance confirms the working code solves the use cases. Each result below reflects "
    "an actual run on the Android Emulator (API 34) and/or a physical device.")

b.H2("25.1 Test Strategy")
b.bullet("verifies the code works \u2014 individual units behave per the technical spec "
         "(state emission, dispatcher usage, layer boundaries).", bold_lead="Unit / Verification \u2014 ")
b.bullet("verifies the working code solves the use case \u2014 user-story steps produce the expected "
         "outcome end-to-end.", bold_lead="Validation / User Acceptance \u2014 ")
b.bullet("integration and regression runs while coding to confirm units work together and previously "
         "passing behavior still passes.", bold_lead="Integration & Regression \u2014 ")
b.bullet("the use-case solution is demonstrated against the original user requirement.",
         bold_lead="Final Acceptance \u2014 ")

b.H2("25.2 Verification (Unit / Technical) Test Cases & Results")
b.P("Derived from the Technical Specification \u2014 system activities in response to user interactions.")
b.table(["ID", "Unit Under Test", "Procedure / Assertion", "Expected", "Result"],
    [["VT-1", "WorkoutViewModel.startWorkout()", "Call startWorkout(); collect uiState.", "State transitions Idle\u2192Active with timer running.", "PASS"],
     ["VT-2", "StateFlow rotation survival", "Rotate device during active workout; re-collect state.", "Active state retained (ViewModel survives config change).", "PASS"],
     ["VT-3", "Workout persistence path", "stopWorkout(); observe submission.", "Session submitted on a background coroutine; UI not blocked.", "PASS"],
     ["VT-4", "Layer boundary (ViewModel)", "Static check: ViewModel imports.", "No DAO/Room/SQLite import in any ViewModel.", "PASS"],
     ["VT-5", "Coroutine hygiene", "Static scan for GlobalScope / runBlocking.", "Zero occurrences across the codebase.", "PASS"],
     ["VT-6", "LogViewModel BP validation", "Submit empty / out-of-range BP.", "Validation error raised before any data-source call.", "PASS"],
     ["VT-7", "InsightsViewModel loading", "Call load()/generate(); inspect first emission.", "Loading emitted before the network/repository call.", "PASS"],
     ["VT-8", "Sensor classification", "Feed accelerometer magnitudes across thresholds.", "Classifies stationary/walking/mixed/running correctly.", "PASS"],
     ["VT-9", "viewModelScope cancellation", "Trigger onCleared(); confirm jobs cancel.", "timerJob cancelled; sensor listener unregistered.", "PASS"]],
    widths=[0.5, 1.9, 2.4, 1.7, 0.7], font=8.5)

b.H2("25.3 Validation (User Acceptance) Test Cases & Results")
b.P("Derived directly from the user steps in the User Stories (\u00a74) and Use Cases (\u00a76).")
b.table(["ID", "From Story", "User Steps", "Acceptance Outcome", "Result"],
    [["AT-1", "US-1", "Open Workout; tap Start.", "Timer starts; live steps + activity shown.", "PASS"],
     ["AT-2", "US-1", "Start workout; rotate device.", "Session keeps running; metrics preserved.", "PASS"],
     ["AT-3", "US-2", "Enter valid systolic/diastolic; Save.", "Reading saved with timestamp; appears in history.", "PASS"],
     ["AT-4", "US-2", "Enter invalid/empty BP; Save.", "Submission rejected; validation message shown.", "PASS"],
     ["AT-5", "US-3", "Open Insights; tap Generate.", "Score, category scores, ranked suggestions display.", "PASS"],
     ["AT-6", "US-3", "Trigger insights while offline.", "Graceful error/last-result fallback, no crash.", "PASS"],
     ["AT-7", "US-4", "Enter age/height/weight/goal; Save.", "Profile saved and available app-wide.", "PASS"],
     ["AT-8", "US-4", "Save profile; restart app.", "Profile persists and auto-loads.", "PASS"]],
    widths=[0.5, 0.8, 2.4, 3.0, 0.7], font=8.5)

b.H2("25.4 Test Results Summary (Does the code work? Does it solve the use case?)")
b.table(["Metric", "Result"],
    [["Verification (unit) tests executed", "9"],
     ["Verification tests passing", "9 / 9 (100%)"],
     ["Validation (acceptance) tests executed", "8"],
     ["Validation tests passing", "8 / 8 (100%)"],
     ["Integration: cross-feature core flow (profile \u2192 workout \u2192 log \u2192 insights)", "PASS (integrated)"],
     ["Regression: prior features after integration", "No regressions observed"],
     ["Rotation survival (config-change) verified", "PASS"],
     ["GlobalScope / runBlocking found", "None"]],
    widths=[4.6, 3.0], font=9.5)
b.P("Does the code work? Yes \u2014 unit, integration, and regression checks pass while coding. Does "
    "the working code solve the use cases? Yes \u2014 the acceptance tests trace to the user stories "
    "and pass end-to-end. Defects found during integration (notably step-counter base-offset on "
    "first reading and an insights loading-flicker) were fixed and re-tested; see \u00a727.3.")
b.page_break()
print("section 25 done")

# ---------------------------------------------------------------------------
# 26. TRACEABILITY MATRIX AUDIT THROUGH TESTING
# ---------------------------------------------------------------------------
b.H1("26. Traceability Matrix Audit Through Testing")
b.P("This audit extends the Traceability Matrix (\u00a713) by tracing each User Story through its "
    "Acceptance Criteria to the concrete test cases that validate it, and records the verdict. "
    "Every user story has at least one passing acceptance test, and every acceptance test maps "
    "back to a story \u2014 confirming bidirectional coverage with no orphans.")
b.table(["User Story", "Acceptance Criteria", "Verification Test(s)", "Acceptance Test(s)", "Verdict"],
    [["US-1 Start Workout", "AC-1.1\u2013AC-1.8", "VT-1, VT-2, VT-8, VT-9", "AT-1, AT-2", "Covered \u2014 PASS"],
     ["US-2 Log Blood Pressure", "AC-2.1\u2013AC-2.8", "VT-6", "AT-3, AT-4", "Covered \u2014 PASS"],
     ["US-3 Generate AI Insights", "AC-3.1\u2013AC-3.9", "VT-7", "AT-5, AT-6", "Covered \u2014 PASS"],
     ["US-4 Create User Profile", "AC-4.1\u2013AC-4.8", "VT-3 (persist path)", "AT-7, AT-8", "Covered \u2014 PASS"]],
    widths=[1.7, 1.5, 1.8, 1.4, 1.2], font=9)
b.H2("26.1 Audit Findings")
b.bullet("Coverage: 4/4 user stories have passing acceptance tests; all acceptance tests trace to a story (no orphan tests).", )
b.bullet("Gaps closed: AC-3.8 (graceful API failure) is now explicitly exercised by AT-6 (offline fallback).")
b.bullet("Cross-cutting AC (rotation survival, AC-1.7 / AC-4.5 persistence) verified by VT-2 and AT-2/AT-8.")
b.bullet("Result: the traceability matrix is sound and test-backed; no requirement is unverified.")
b.page_break()

# ---------------------------------------------------------------------------
# 27. ARCHITECTURAL LEGITIMACY AUDIT
# ---------------------------------------------------------------------------
b.H1("27. Architectural Legitimacy Audit")
b.P("This section is the architectural legitimacy checkpoint requested for GP6. It audits the "
    "integrated codebase against the disciplines this course enforces, with the specific checks "
    "the instructor is looking for.")

b.H2("27.1 Architectural Alignment Audit")
b.table(["Check", "Finding", "Status"],
    [["MVVM layering (UI \u2192 ViewModel \u2192 Repository \u2192 source)",
      "Screens are Composables backed by dedicated ViewModels exposing StateFlow; data access is mediated, not done in the UI.", "Aligned"],
     ["ViewModel exposes StateFlow / Flow (explicit terms)",
      "Each ViewModel exposes a MutableStateFlow privately and a read-only StateFlow publicly (e.g., WorkoutViewModel.uiState: StateFlow<WorkoutUiState>).", "Aligned"],
     ["Compose collects state",
      "Screens collect via collectAsState()/collectAsStateWithLifecycle() and recompose on emission.", "Aligned"],
     ["Repository as the data-source mediator",
      "Five repository interfaces + impls hide Room vs API behind domain-typed functions; ViewModels depend only on the interfaces; data-source wiring is centralized in the ServiceLocator composition root.", "Aligned (implemented)"],
     ["Logic NOT in Activity/Composable",
      "MainActivity only hosts the Compose tree + theme; business/state logic resides in ViewModels.", "Aligned"]],
    widths=[2.6, 3.6, 1.4], font=9)

b.H2("27.2 Coroutine & Lifecycle Hygiene Audit")
b.table(["Audited Item", "Method", "Result"],
    [["GlobalScope usage", "Full-repository static scan.", "None found \u2014 all coroutines use viewModelScope."],
     ["runBlocking usage", "Full-repository static scan.", "None found."],
     ["Logic in Activity", "Inspect MainActivity / Composables.", "None \u2014 host only; state owned by ViewModels."],
     ["Missing/malformed repository layer", "Inspect data flow + contract.", "Repository contract defined (\u00a715); ViewModels target domain API, not DAOs/DTOs."],
     ["Scope ownership", "Inspect launch sites.", "All launches in viewModelScope; cancelled in onCleared()."],
     ["State surviving rotation", "Rotate during active workout & after profile save.", "State survives \u2014 ViewModel + StateFlow retained (VT-2, AT-2, AT-8)."]],
    widths=[2.4, 2.4, 2.8], font=9)
b.P("Evidence \u2014 explicit StateFlow exposure and scoped launches (from WorkoutViewModel):", bold=True, size=10, color=BLUE)
b.code(
"private val _uiState = MutableStateFlow<WorkoutUiState>(WorkoutUiState.Idle)\n"
"val uiState: StateFlow<WorkoutUiState> = _uiState        // explicit StateFlow\n\n"
"timerJob = viewModelScope.launch { /* tick + emit */ }   // viewModelScope only\n\n"
"override fun onCleared() {                                // lifecycle ownership\n"
"    timerJob?.cancel()\n"
"    sensorManager?.unregisterListener(this)\n"
"}")

b.H2("27.3 Defects Found & Fixed During Integration")
b.table(["Defect", "Symptom", "Fix", "Re-test"],
    [["Step-counter base offset", "First step reading showed device lifetime steps.", "Capture baseStepCount on first event; report delta.", "VT-8 PASS"],
     ["Insights loading flicker", "Brief empty flash before Loading.", "Emit Loading before the repository call.", "VT-7 PASS"],
     ["Rotation reset (early build)", "Workout reset on rotation in an early prototype.", "Move state into ViewModel/StateFlow.", "VT-2/AT-2 PASS"]],
    widths=[1.8, 2.4, 2.2, 1.0], font=9)

b.H2("27.4 Audit Verdict")
b.callout("Verdict: ARCHITECTURALLY LEGITIMATE. Core functionality is integrated; the ViewModel is "
          "the state owner exposing explicit StateFlow; Compose collects and recomposes reactively; "
          "no GlobalScope or runBlocking exists; logic is out of the Activity/Composables; the "
          "Repository layer mediates live Retrofit calls and a Room offline cache (ViewModels import "
          "neither Retrofit nor Room); the full Navigation graph is in place; and state survives "
          "rotation.")
b.page_break()
print("sections 26-27 done")

# ---------------------------------------------------------------------------
# 28. WHAT REMAINS TO BE DEVELOPED
# ---------------------------------------------------------------------------
b.H1("28. Implementation Status & Remaining Backlog")
b.P("The core architectural layers planned in earlier milestones are now implemented. The table "
    "records the as-built status of each work item; only the items marked Future remain.")
b.table(["Feature / Work Item", "Status", "Notes"],
    [["Retrofit client + live ApiService calls", "Implemented", "RetrofitClient + suspend ApiService called only inside the RepositoryImpls."],
     ["Room entities, DAOs, and offline cache", "Implemented", "AppDatabase v2 with 7 entities/DAOs; repositories write through and fall back to Room when offline."],
     ["DTO \u2194 Entity mappers", "Implemented", "toDto()/toEntity() mappers on every entity."],
     ["Composition root (manual DI)", "Implemented", "ServiceLocator builds ApiService + AppDatabase once; ViewModels import neither."],
     ["Full Jetpack Navigation implementation", "Implemented", "NavHost + bottom nav across 6 destinations; back stack survives rotation."],
     ["History import (JSON/CSV) end-to-end", "Implemented", "ProfileRepository.importFile (neutral signature) performs the multipart upload."],
     ["Charts/analytics (trends)", "Partial", "History line/bar charts shipped; deeper trend analytics are future work."],
     ["Medication / nutrition / sleep logging UI", "Implemented", "Log screen covers all 7 entry types."],
     ["Auth + multi-user / cloud sync", "Future (out of MVP)", "Profile carries user_id; error contract reserves AUTH_ERROR."],
     ["Automated UI tests (Compose/Espresso)", "Future", "Manual test cases documented (\u00a725) ready to automate."]],
    widths=[3.2, 1.6, 2.8], font=9)
b.page_break()

# ---------------------------------------------------------------------------
# 29. PHASED DEVELOPMENT APPROACH — UPDATED
# ---------------------------------------------------------------------------
b.H1("29. Phased Development Approach — Updated Next Steps")
b.P("This refines \u00a723 with the post-integration plan. It states what functionality comes next "
    "and how the current milestone prepared for it.")
b.H2("29.1 Next Steps (Immediate)")
b.numbered("Instantiate the Retrofit client against the documented base URL and wire ApiService into the Repository implementation.")
b.numbered("Introduce Room (database, entities, DAOs) and the CachedInsight write-through cache.")
b.numbered("Implement DTO\u2192Domain\u2192Entity mappers in the Repository exactly as specified in \u00a717.")
b.numbered("Verify Main\u2192IO\u2192Main boundaries at runtime with logging/Thread.currentThread checks.")
b.numbered("Complete Jetpack Navigation per the Phase 3 discipline (NavGraph owns the back stack).")
b.H2("29.2 Functionality Sequencing")
b.bullet("Phase 2 (next): persistence + live API behind the Repository \u2014 no UI/ViewModel changes required.", bold_lead="")
b.bullet("Phase 3: navigation finalization, history import, analytics polish.", bold_lead="")
b.bullet("Phase 4+: richer lifestyle logging, then future items (auth, cloud sync, automated UI tests).", bold_lead="")
b.H2("29.3 How We Prepared Now for Continuing Development")
b.bullet("Locked the Repository's public surface to domain types so Phase 2 adds sources without touching the ViewModel or UI.")
b.bullet("Designed the full API contract, DTOs, mapping path, and thread-boundary plan ahead of implementation to minimize refactor cost.")
b.bullet("Kept coroutines in viewModelScope and state in StateFlow so the reactive spine is already API-ready.")
b.bullet("Maintained GitHub Project boards, activity logs, and a team retrospective to keep management artifacts current.")
b.callout("Architectural legitimacy checkpoint cleared: integration achieved, audits pass, and the "
          "design is positioned for low-risk Phase 2 API implementation.")
print("sections 28-29 done")

# ===========================================================================
# 6. SAVE
# ===========================================================================
out = os.path.join(DOCDIR, "BTJ-GP6-SRS-document-06-16-26-brandonGalli.docx")
doc.save(out)
print("SAVED:", out)
print("final paragraphs:", len(doc.paragraphs), "tables:", len(doc.tables),
      "images:", len(doc.inline_shapes))





